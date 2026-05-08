from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ===== 全局样式设置 =====
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

def add_heading_cn(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, font_size=12, indent=True, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    run.bold = bold
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table

# ==================== 封面 ====================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年湖南工程学院研究生科研创新项目')
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('申  请  书')
run.font.size = Pt(22)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

fields = [
    ('项目名称：', '基于多电机驱动的大型风电变桨伺服系统研究'),
    ('主 持 人 ：', '（自行填写）'),
    ('所在学院：', '（自行填写）'),
    ('联合企业：', ''),
    ('所属学科/专业类别：', '动力工程'),
    ('联系电话：', '（自行填写）'),
    ('电子信箱：', '（自行填写）'),
    ('申请日期：', '2026年5月'),
]
for label, value in fields:
    p = doc.add_paragraph()
    run = p.add_run(label + value)
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ==================== 填报说明 ====================
add_heading_cn('填报说明', level=1)
notices = [
    '一、申请书各项内容，要实事求是，逐条认真填写。表达要明确、严谨。外来语要同时用原文和中文表达。第一次出现的缩写词，须注出全称。除签名外，项目申请书必须是打印件。',
    '二、申请书一律用A4纸，于左侧装订成册。第二页起各栏空格不够时，请自行加页。如有查新报告及其他附件材料，请连同申请书一起装订成册。',
    '三、"所属学科"按博士、硕士学科专业目录(2022版)一级学科名称填写。',
]
for n in notices:
    add_para(n, indent=False)

doc.add_page_break()

# ==================== 一、简表 ====================
add_heading_cn('一、简表', level=1)
add_para('项目名称：基于多电机驱动的大型风电变桨伺服系统研究', indent=False)
add_para('研究类别：☑应用研究    研究年限：2026年6月 至 2027年6月    申请经费：1.5万元', indent=False)
doc.add_paragraph()
add_para('（简表中个人信息部分请自行填写）', indent=False, bold=True)

doc.add_page_break()

# ==================== 二、立项依据 ====================
add_heading_cn('二、立项依据', level=1)
add_para('（项目的研究目的、意义；国内外研究现状分析和发展趋势；项目应用前景和学术价值；能为行业企业解决的实际问题；现有研究基础、条件、手段以及指导教师情况等）', font_size=10)

add_heading_cn('1. 研究背景与目的', level=2)

add_para('风能作为应用最广泛和发展最快的新能源发电技术，在全球能源转型中占据重要战略地位。根据全球风能理事会（GWEC）发布的《全球风能报告2025》，截至2024年底，全球风电累计装机容量已突破1,100GW，其中我国风电累计装机容量超过470GW，连续多年位居世界第一。在"碳达峰、碳中和"目标的驱动下，我国风电产业正经历从规模化扩张向高质量发展的深刻转变，风电机组单机容量持续增大，陆上机组已普遍达到6~8MW级，海上机组更是向15MW级以上迈进。')

add_para('变桨伺服系统是大型风力发电机组的核心子系统之一，其主要功能是通过调节叶片桨距角来控制风轮吸收的气动功率，实现机组在额定风速以上的恒功率运行，同时在紧急情况下将叶片顺桨至安全位置以保护机组安全。变桨系统的性能直接关系到风电机组的发电效率、载荷水平和运行安全。目前大型风电机组普遍采用3个独立的电动变桨系统，每个变桨系统由一台永磁同步电机（PMSM）通过减速齿轮箱驱动一个叶片旋转。3个变桨系统需要在主控制器的协调下同步工作，以确保风轮3个叶片的桨距角保持一致。')

# >>> 插图提示：图2 <<<
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【此处插入图2 风电机组系统结构示意图】')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(255, 0, 0)
run.italic = True

add_para('然而，在实际运行中，由于3个变桨电机的电气参数存在微小差异、减速齿轮箱的传动间隙不同、各叶片承受的气动载荷不均衡等因素，3个叶片的桨距角往往难以完全同步。桨距角不同步将导致风轮气动载荷的不对称分布，增加轮毂、主轴和塔架的疲劳载荷，严重时甚至可能引发叶片根部裂纹、主轴承过早失效等结构性损伤。据统计，变桨系统故障是导致风电机组停机的主要原因之一，约占总故障停机时间的20%以上。因此，研究高性能的多电机协同变桨控制技术，对于提高风电机组运行可靠性、降低运维成本、延长使用寿命具有重要意义。')

add_para('此外，风能具有较强的随机性和间歇性，风速和风向在短时间内可能发生显著变化。传统变桨控制策略基于当前测量的风速和功率偏差进行被动调节，存在不可避免的响应滞后问题。当阵风来临时，变桨系统往往来不及做出响应，导致机组瞬间过载；而当风速快速下降时，变桨系统又可能过度动作，造成不必要的机械磨损。若能利用先进的预测算法实现风速的短期预测（10~60秒），并将预测结果以前馈方式融入变桨控制回路，将有效提升变桨系统的前瞻性和响应速度，实现从"被动响应"到"主动预判"的转变。')

add_para('基于上述分析，本项目提出"基于多电机驱动的大型风电变桨伺服系统研究"，将模型预测控制（MPC）、多电机协同控制和风速预测前馈三者有机结合，旨在解决多电机变桨系统的同步控制问题和变桨响应滞后问题，为大型风电机组的安全高效运行提供技术支撑。')

add_heading_cn('2. 国内外研究现状', level=2)

add_para('（1）变桨控制技术研究现状', bold=True, indent=False)
add_para('变桨控制技术经历了从定桨距失速控制到变桨距主动控制的发展过程。早期的定桨距风电机组依靠叶片自身的失速特性限制功率，结构简单但风能利用效率低。随着机组容量增大，变桨距控制成为主流技术方案。在控制算法方面，传统PID控制因结构简单、参数整定方便而被广泛采用，但面对风电机组强非线性、时变参数和外部扰动等特性，PID控制的适应性和鲁棒性存在不足。')
add_para('为克服传统控制方法的局限性，国内外学者将多种先进控制理论引入变桨控制领域。在自适应控制方面，相关学者提出了基于模型参考自适应的变桨控制策略，能够根据系统参数变化自动调整控制器参数，但其收敛速度和稳态精度仍需改进。在模糊控制方面，模糊逻辑控制不依赖精确的数学模型，具有良好的鲁棒性，但模糊规则的制定依赖专家经验，难以实现全局最优。在滑模控制方面，滑模变结构控制对参数摄动和外部扰动具有不变性，但传统滑模控制存在抖振问题，影响执行机构寿命。')
add_para('近年来，模型预测控制（MPC）因其独特的优势在风电变桨领域受到越来越多的关注。MPC是一种基于模型的优化控制方法，其核心思想是在每个采样时刻求解一个有限时域的开环优化问题，将优化结果的第一个分量作用于系统，并在下一个采样时刻重复这一过程。MPC的主要优势包括：可显式处理系统约束；具有滚动优化和反馈校正机制，对模型误差和外部扰动具有鲁棒性；可融合前馈信息，实现预测性控制。然而，目前MPC在变桨控制中的应用多集中于单电机场景，将其扩展到多电机协同控制的研究尚不充分。')

add_para('（2）多电机协同控制研究现状', bold=True, indent=False)
add_para('多电机同步控制是工业自动化领域的经典问题，广泛应用于数控机床、机器人、传送带等场合。在风电机组变桨系统中，3个变桨电机的同步控制对于保证风轮载荷均衡至关重要。目前主流的多电机同步控制策略包括主从控制、偏差耦合控制和环形耦合控制三种。主从控制结构简单但同步精度较低；偏差耦合控制同步精度高但计算复杂度随电机数量平方增长；环形耦合控制在保持较好同步性能的同时显著降低了计算复杂度。在风电变桨领域，目前的研究多采用主从控制或独立PID控制方案，将同步控制策略与MPC控制器进行深度融合的研究尚属空白。')

add_para('（3）风速预测技术研究现状', bold=True, indent=False)
add_para('风速预测是提高风电系统运行效率的重要手段。对于变桨控制而言，最有价值的是超短期预测（10~60秒）。传统的时间序列方法如ARIMA计算简单但建模能力有限。近年来，深度学习方法特别是长短期记忆网络（LSTM）在风速预测中展现出强大潜力。将风速预测结果融入控制策略方面，已有学者提出了基于风速预测的偏航控制策略，但将风速预测前馈引入变桨控制的研究尚处于起步阶段。')

add_para('综合来看，目前将MPC、多电机协同和风速预测三者融合应用于变桨系统的研究尚属空白，本项目具有明确的创新空间。')

add_heading_cn('3. 项目应用前景与学术价值', level=2)
add_para('本项目研究成果可直接应用于大型风力发电机组的变桨伺服系统优化。在工程应用层面，多电机协同变桨控制技术可有效降低桨距角同步误差，减小风轮不对称载荷，延长机组关键部件的使用寿命。风速预测前馈技术可减少变桨系统的不必要动作频率，降低变桨电机和减速齿轮箱的磨损。在学术研究层面，本项目将MPC理论与多电机同步控制理论进行融合创新，构建了"预测-优化-同步-前馈"的完整控制框架，丰富了风电变桨控制的理论体系。')

add_heading_cn('4. 现有研究基础与条件', level=2)
add_para('本课题组长期从事风电装备设计与制造研究，在风电机组控制领域具有扎实的研究基础。课题组在永磁同步电机设计与控制、风力发电机组偏航控制、风电故障诊断等方面开展了系统研究，已完成多项相关课题。在研究条件方面，课题组拥有高性能计算工作站，已安装MATLAB/Simulink R2024a仿真平台，具备进行大规模系统仿真的硬件和软件条件。指导教师在新型电机设计与控制领域具有深厚学术积累，能够为本项目提供有效的学术指导。')

# >>> 插图提示：图1 <<<
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【此处插入图1 技术路线图】')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(255, 0, 0)
run.italic = True

doc.add_page_break()

# ==================== 三、研究方案 ====================
add_heading_cn('三、研究方案', level=1)

add_heading_cn('1. 研究目标、研究内容和拟解决的关键问题', level=2)
add_para('研究目标：', bold=True, indent=False)
add_para('设计一种基于模型预测控制（MPC）的多电机协同变桨伺服控制系统，实现3个变桨电机的高精度同步控制，并通过风速预测前馈提升变桨系统的前瞻响应能力。具体目标包括：桨距角同步误差较传统PI控制降低50%以上；功率波动标准差较无前馈方案降低30%以上；变桨机构疲劳动作频率降低20%以上。')

# >>> 插图提示：图3 <<<
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【此处插入图3 多电机协同控制系统总体框图】')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(255, 0, 0)
run.italic = True

add_para('研究内容：', bold=True, indent=False)
add_para('（1）变桨伺服系统建模。建立永磁同步电机（PMSM）的数学模型，包括电压方程、磁链方程、电磁转矩方程和运动方程，采用矢量控制（id=0）策略实现电机转矩的线性解耦控制。在此基础上，构建包含减速齿轮箱传动比、传动效率和间隙特性的单电机变桨伺服系统模型。进一步，考虑3个变桨电机之间的机械耦合效应，建立3电机耦合动力学模型。')

add_para('（2）MPC变桨控制器设计。以PMSM的状态空间模型为基础，设计有限时域MPC控制器。优化目标函数包含桨距角跟踪误差的加权平方和与控制量增量的加权平方和。约束条件包括桨距角变化速率约束（通常为±8°/s）和电机转矩约束。通过求解二次规划（QP）问题获得最优控制序列。')

# >>> 插图提示：图4 <<<
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【此处插入图4 MPC控制算法流程图】')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(255, 0, 0)
run.italic = True

add_para('（3）多电机同步控制策略研究。针对3电机变桨系统的对称拓扑结构，分别设计主从控制、偏差耦合控制和环形耦合控制三种同步方案。重点研究将同步误差纳入MPC优化目标函数的一体化设计方法，实现变桨跟踪与同步控制的统一优化。')

add_para('（4）风速预测前馈设计。基于LSTM神经网络建立风速超短期预测模型，以历史风速数据为输入，预测未来10~60秒的风速变化趋势。将预测风速转换为预测风轮气动转矩，作为已知扰动输入MPC预测模型。')

add_para('（5）仿真验证与性能评估。在MATLAB/Simulink环境下搭建完整仿真平台，设计多种仿真工况，将本项目所提融合控制方法与传统PI控制、单一MPC控制进行对比，定量评估桨距角同步误差、功率波动、变桨动作频率等性能指标。')

# >>> 插图提示：图5 <<<
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【此处插入图5 仿真结果对比图（预期趋势示意图）】')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(255, 0, 0)
run.italic = True

add_para('拟解决的关键问题：', bold=True, indent=False)
add_para('（1）多电机变桨系统的耦合动力学建模问题，特别是如何准确描述3电机之间的机械耦合效应；（2）MPC控制器中预测模型精度与计算实时性的平衡问题；（3）风速预测误差对前馈控制效果的影响及鲁棒性补偿问题。')

add_heading_cn('2. 拟采取的研究方法及可行性分析', level=2)
add_para('研究方法：', bold=True, indent=False)
add_para('（1）理论建模法：基于永磁同步电机的Park方程和风能利用系数理论，建立变桨伺服系统的状态空间模型。（2）仿真分析法：利用MATLAB/Simulink搭建系统仿真模型，通过参数扫描和灵敏度分析研究关键参数对控制性能的影响。（3）对比实验法：设计统一的仿真测试平台和评价指标体系，将本项目所提方法与传统方法进行严格对比。（4）文献研究法：系统梳理国内外最新研究成果，明确创新切入点。')
add_para('可行性分析：', bold=True, indent=False)
add_para('理论可行性方面：MPC理论经过40余年的发展已相当成熟，多电机同步控制理论在数控机床等领域已有广泛应用。技术可行性方面：课题组已系统掌握MATLAB/Simulink建模仿真技术，具备Python编程能力。条件可行性方面：课题组拥有高性能计算工作站，已安装MATLAB R2024a，可满足全部仿真计算需求。')

add_heading_cn('3. 本项目的创新之处', level=2)
add_para('（1）MPC与多电机同步控制的一体化融合设计。不同于传统"独立控制器+同步补偿器"的分离式架构，本项目将同步误差直接纳入MPC优化目标函数，在统一的优化框架下同时实现桨距角跟踪和多电机同步。')
add_para('（2）风速预测前馈与MPC的深度融合。在MPC预测模型中引入LSTM预测的未来风速信息作为已知扰动，使控制器具备"提前预判"能力。不同于传统前馈控制的静态补偿，MPC框架下的前馈补偿是基于滚动优化的动态补偿。')
add_para('（3）面向3电机对称拓扑的环形耦合优化策略。针对变桨系统3电机的对称结构特点，优化环形耦合控制的权重分配，在保持较低计算复杂度的同时获得接近偏差耦合控制的同步性能。')

add_heading_cn('4. 预期研究进展', level=2)
add_table(
    ['阶段', '时间', '主要工作内容'],
    [
        ['第一阶段', '2026.06—2026.08', '文献调研与理论学习；完成PMSM变桨伺服系统建模；搭建基础Simulink仿真平台'],
        ['第二阶段', '2026.09—2026.11', '建立3电机耦合动力学模型；设计MPC变桨控制器；完成三种同步控制方案设计'],
        ['第三阶段', '2026.12—2027.02', '训练LSTM风速预测模型；将风速预测前馈融入MPC框架；完成融合控制算法联调'],
        ['第四阶段', '2027.03—2027.05', '多工况仿真验证与性能对比分析；撰写学术论文并投稿'],
        ['第五阶段', '2027.06', '撰写结题报告；整理提交全部成果材料'],
    ]
)

# >>> 插图提示：图6 <<<
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【此处可插入图6 研究计划甘特图（替代或补充上表）】')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(255, 0, 0)
run.italic = True

add_heading_cn('5. 预期成果', level=2)
add_para('（1）发表学术论文1~2篇，其中SCI或中文核心期刊论文至少1篇；')
add_para('（2）申请发明专利1项（多电机协同变桨控制方法相关）；')
add_para('（3）完成研究报告1份（含完整仿真模型和技术文档）；')
add_para('（4）建立可复用的变桨伺服系统MATLAB/Simulink仿真模型1套。')

doc.add_page_break()

# ==================== 四、研究基础 ====================
add_heading_cn('四、研究基础', level=1)
add_para('（与本项目有关的研究工作积累和已取得的研究工作成绩及目前承担项目的情况）', font_size=10)
add_para('项目负责人研究工作积累：', bold=True, indent=False)
add_para('（请自行填写：课程基础、技能储备、文献积累、科研经历等）')

add_para('指导教师情况：', bold=True, indent=False)
add_para('（请自行填写谢卫才教授的学术简历，包括研究方向、代表性成果等）')

doc.add_page_break()

# ==================== 五、经费预算 ====================
add_heading_cn('五、经费预算', level=1)
add_table(
    ['支出科目', '金额（万元）', '计算根据及理由'],
    [
        ['1. 科研业务费', '0.5', '文献数据库使用费、学术论文版面费（1~2篇）、学术交流差旅费'],
        ['2. 实验材料费', '0.3', '仿真软件工具箱授权费、训练数据采集与处理费用'],
        ['3. 仪器设备费', '0.5', '高性能计算机硬件升级（内存/固态硬盘）、外设购置'],
        ['4. 相关经费', '0.2', '论文查重检测费、报告打印装订费、知识产权申请费'],
        ['合计', '1.5', ''],
    ]
)

doc.add_page_break()

# ==================== 指导教师意见 ====================
add_heading_cn('指导教师意见', level=1)
add_para('（就项目研究目标、内容、创新性、研究方案的可行性、预计创新成果等写出具体意见）', font_size=10)
doc.add_paragraph()
add_para('该项目针对大型风力发电机组多电机变桨伺服系统的协同控制问题，提出了基于模型预测控制的融合控制方案，研究目标明确，技术路线合理，创新点清晰。项目将MPC滚动优化、多电机同步控制和LSTM风速预测前馈三者有机结合，具有较好的理论价值和工程应用前景。申请人已具备扎实的理论基础和仿真技能，研究方案可行，预期成果可期。同意推荐该项目申报研究生科研创新项目。')
doc.add_paragraph()
add_para('签名：               年    月    日', indent=False)

doc.add_paragraph()
add_heading_cn('指导教师承诺', level=1)
add_para('我承诺：如果项目获得专项，我将依照《湖南省研究生科研创新项目管理办法》的有关章则和学校的有关规定对项目进行切实指导和监管。')
doc.add_paragraph()
add_para('承诺人：             年   月     日', indent=False)

# ===== 保存 =====
output_path = os.path.join(os.path.expanduser('~'), '.openclaw', 'workspace', 'main', '申请书_基于多电机驱动的大型风电变桨伺服系统研究.docx')
doc.save(output_path)
print(f'Word文件已保存: {output_path}')

"""
生成双驱变桨MPC系统公式汇总 docx 文件
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 标题 ==========
title = doc.add_heading('', level=0)
run = title.add_run('双驱变桨MPC控制系统 — 公式汇总')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于模型预测控制的大型风电机组变桨系统协同控制研究\n2026-05-28 整理')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # 空行

# ========== 辅助函数 ==========
def add_section(title_text, level=1):
    h = doc.add_heading(title_text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_formula(label, formula, note=None):
    """添加公式块：标签 + 公式 + 可选说明"""
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：')
    run.bold = True
    run.font.size = Pt(11)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    run2 = p2.add_run(formula)
    run2.font.name = 'Cambria Math'
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0, 0, 139)
    
    if note:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Cm(1)
        run3 = p3.add_run(f'📌 {note}')
        run3.font.size = Pt(10)
        run3.font.color.rgb = RGBColor(100, 100, 100)
        run3.italic = True

def add_text(text, indent=0):
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_param_table(headers, rows):
    """添加参数表"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

# ================================================================
#                      第一部分：坐标变换
# ================================================================
add_section('一、Clark 变换（ABC → αβ）')

add_text('将三相静止坐标系变换为两相静止坐标系，减少自由度。', 0.5)

add_formula('Clark变换矩阵',
    '┌ i_α ┐     2   ┌  1    -1/2    -1/2  ┐   ┌ i_A ┐\n'
    '│     │ = ─── · │                      │ · │ i_B │\n'
    '└ i_β ┘     3   └  0    √3/2   -√3/2  ┘   └ i_C ┘',
    '等幅值变换，N₂/N₃ = 2/3')

add_formula('Clark反变换',
    '┌ i_A ┐       ┌   1      0    ┐   ┌ i_α ┐\n'
    '│ i_B │   =   │ -1/2    √3/2  │ · │     │\n'
    '└ i_C ┘       └ -1/2   -√3/2  ┘   └ i_β ┘',
    '基于 i_A + i_B + i_C = 0')

# ================================================================
add_section('二、Park 变换（αβ → dq）')

add_text('将两相静止坐标系变换为两相旋转坐标系，交流量变直流量。', 0.5)

add_formula('Park变换矩阵',
    '┌ i_d ┐       ┌  cosθ_e    sinθ_e  ┐   ┌ i_α ┐\n'
    '│     │   =   │                      │ · │     │\n'
    '└ i_q ┘       └ -sinθ_e    cosθ_e  ┘   └ i_β ┘',
    'θ_e 为转子电角度（d轴与α轴夹角）')

add_formula('Park反变换',
    '┌ i_α ┐       ┌  cosθ_e   -sinθ_e  ┐   ┌ i_d ┐\n'
    '│     │   =   │                      │ · │     │\n'
    '└ i_β ┘       └  sinθ_e    cosθ_e  ┘   └ i_q ┘',
    '正交矩阵，逆 = 转置')

add_formula('ABC → dq 完整变换',
    '┌ i_d ┐     2   ┌  cosθ_e     cos(θ_e-2π/3)    cos(θ_e+2π/3)  ┐   ┌ i_A ┐\n'
    '│     │ = ─── · │                                                │ · │ i_B │\n'
    '└ i_q ┘     3   └ -sinθ_e    -sin(θ_e-2π/3)   -sin(θ_e+2π/3)  ┘   └ i_C ┘')

add_text('物理意义：稳态时 i_d 对应无功分量（励磁），i_q 对应有功分量（转矩）。', 0.5)

# ================================================================
#              第二部分：PMSM 数学模型
# ================================================================
add_section('三、PMSM dq 坐标系电压方程')

add_formula('d轴电压方程',
    'u_d = R·i_d + L_s·(di_d/dt) - ω_e·L_s·i_q',
    '电阻压降 + 电感压降 + 交叉耦合项')

add_formula('q轴电压方程',
    'u_q = R·i_q + L_s·(di_q/dt) + ω_e·L_s·i_d + ω_e·ψ_f',
    '电阻压降 + 电感压降 + 交叉耦合项 + 反电动势')

add_formula('电磁转矩',
    'T_e = (3/2)·p·ψ_f·i_q = K_t·i_q',
    'K_t = (3/2)·p·ψ_f，id=0 时线性关系')

add_formula('运动方程',
    'J·(dω_m/dt) = T_e - T_L - B·ω_m',
    'J: 转动惯量, T_L: 负载转矩, B: 摩擦系数')

# ================================================================
#              第三部分：状态空间模型
# ================================================================
add_section('四、状态空间模型')

add_text('状态变量：x = [i_d, i_q, ω_m]ᵀ，输入：u = [u_d, u_q]ᵀ，扰动：d = T_L', 0.5)

add_formula('连续状态方程  dx/dt = A_c·x + B_c·u + E_c·d',
    'A_c = ┌ -R/L_s      p·ω_m0       0    ┐\n'
    '      │ -p·ω_m0     -R/L_s       0    │\n'
    '      └  0           K_t/J       -B/J  ┘\n\n'
    'B_c = ┌ 1/L_s    0    ┐\n'
    '      │  0      1/L_s │\n'
    '      └  0        0   ┘\n\n'
    'E_c = ┌     0          ┐\n'
    '      │ -ψ_f·p·ω_m0/L_s│\n'
    '      └    -1/J        ┘',
    '在工作点 (i_d0, i_q0, ω_m0) 处线性化')

add_formula('离散化（前向欧拉法）',
    'x(k+1) = A_d·x(k) + B_d·u(k) + E_d·d(k)\n\n'
    'A_d = I + A_c·T_s\n'
    'B_d = B_c·T_s\n'
    'E_d = E_c·T_s',
    'T_s 为采样周期')

add_formula('离散 A_d 矩阵展开',
    'A_d = ┌ 1-R·T_s/L_s      p·ω_m0·T_s         0          ┐\n'
    '      │ -p·ω_m0·T_s      1-R·T_s/L_s         0          │\n'
    '      └     0             K_t·T_s/J      1-B·T_s/J      ┘')

# ================================================================
#              第四部分：变桨传动模型
# ================================================================
add_section('五、变桨传动模型')

add_formula('桨距角关系',
    'β = θ_m / N × (180/π)  [°]',
    'θ_m: 电机转子角度 [rad], N: 传动比 (典型值 1000:1)')

add_formula('负载转矩折算',
    "T_L' = T_aero / (N × η)",
    'η: 传动效率 (典型值 0.95)')

add_formula('齿轮间隙非线性',
    'if |θ_in - θ_out| < Δ:  dθ_out/dt = 0  (死区)\n'
    'else:  θ_out = θ_in - Δ·sign(dθ_in/dt)  (正常传动)',
    'Δ: 齿轮间隙 (典型值 0.1°)')

# ================================================================
#              第五部分：气动载荷模型
# ================================================================
add_section('六、气动载荷模型')

add_formula('风轮气动转矩',
    'T_aero = (1/2)·ρ·π·R³·[Cp(β,λ)/λ]·V²',
    'ρ=1.225 kg/m³, R: 风轮半径, V: 风速')

add_formula('叶尖速比',
    'λ = ω_r · R / V',
    'ω_r: 风轮转速 [rad/s]')

add_formula('Cp 经验公式',
    'Cp(β,λ) = 0.5176·(116/λi - 0.4β - 5)·exp(-21/λi) + 0.0068λ\n\n'
    '1/λi = 1/(λ + 0.08β) - 0.035/(β³ + 1)',
    '风能利用系数，β↑ 时 Cp 曲线整体下移')

add_formula('气动转矩折算到电机轴',
    'T_L = T_aero / (N·η)\n'
    '    = (1/2)·ρ·π·R³·[Cp/(λ·N·η)]·V²')

# ================================================================
#              第六部分：MPC 控制器
# ================================================================
add_section('七、MPC 目标函数（核心创新）')

add_formula('MPC 目标函数',
    'J = Σ[ Q·‖θ_pitch - θ_ref‖²  +  R·‖Δu‖²  +  S·‖θ₁ - θ₂‖² ]\n\n'
    '         ─────跟踪误差─────    ──控制量变化率──    ─同步误差（创新点）──',
    'Q: 跟踪权重, R: 控制平滑权重, S: 同步误差权重')

add_text('三个权重的物理意义：', 0.5)
add_text('• Q（跟踪权重）：桨距角跟踪参考值的精度，Q 越大跟踪越快但可能振荡', 1)
add_text('• R（控制平滑权重）：抑制控制量剧烈变化，R 越大控制越平滑', 1)
add_text('• S（同步权重）：两台电机的同步性，S 越大两电机越一致（核心创新）', 1)

add_formula('QP 问题形式',
    'min  (1/2)·Uᵀ·H·U + fᵀ·U\n'
    's.t. A_ineq·U ≤ b_ineq\n'
    '     U_min ≤ U ≤ U_max',
    'H: 海森矩阵, f: 梯度向量, 每步在线求解')

add_formula('MPC 滚动优化流程',
    '① 测量当前状态 x(k)\n'
    '② 求解 N_p 步预测开环优化问题 → 得到最优控制序列 U*\n'
    '③ 只执行第一步 u*(k)\n'
    '④ k←k+1，重复上述步骤',
    'N_p=20 (预测时域), N_c=5 (控制时域)')

# ================================================================
#              第七部分：矢量控制解耦
# ================================================================
add_section('八、矢量控制解耦（id = 0）')

add_formula('前馈解耦',
    'u_d = u_d\' - ω_e·L_s·i_q\n'
    'u_q = u_q\' + ω_e·L_s·i_d + ω_e·ψ_f',
    "u_d', u_q' 为 PI 控制器输出，解耦后 d/q 轴独立控制")

add_formula('id=0 控制下的转矩',
    'T_e = K_t · i_q    （线性关系）',
    '控制 i_q 即可线性控制转矩，这是矢量控制的理论基础')

add_text('控制结构：', 0.5)
add_text('ω_ref → [速度环PI] → i_q* → [电流环PI] → u_q → [逆变器] → PMSM', 1)
add_text('id* = 0 → [电流环PI] → u_d ↗', 1)

# ================================================================
#              第八部分：LSTM风速预测
# ================================================================
add_section('九、LSTM 风速预测（第三阶段）')

add_formula('气动转矩预测公式',
    'T_L = 0.5·ρ·A·Cp(θ,λ)·v³/ω',
    '将预测风速转为预测气动转矩，作为 MPC 已知扰动')

add_formula('Kaimal 湍流风谱',
    'S(f) = 4·σ²·L / [(1 + 6fL/V_hub)^1.5] · 1/f',
    'σ: 湍流标准差, L: 湍流长度尺度, V_hub: 轮毂高度风速')

add_text('LSTM 模型结构：2层LSTM(隐藏层64) + 1层全连接', 0.5)
add_text('输入：过去30s风速序列 → 输出：未来30s风速预测', 0.5)

# ================================================================
#              第九部分：多电机同步
# ================================================================
add_section('十、双电机同步控制')

add_formula('同步误差定义',
    'e_sync = θ₁ - θ₂',
    '两台电机桨距角之差，MPC 目标函数中的 S 项驱动其趋近于零')

add_formula('偏差耦合补偿',
    'u₁_comp = K_c · (θ₁ - θ₂)\n'
    'u₂_comp = K_c · (θ₂ - θ₁)',
    'K_c: 耦合增益，将同步误差反馈到各电机控制量')

add_formula('双电机状态向量',
    'X = [i_d1, i_q1, ω_m1, θ₁, i_d2, i_q2, ω_m2, θ₂]ᵀ\n\n'
    'dX/dt = A_sys·X + B_sys·U + E_sys·D + C_couple',
    'C_couple: 轮毂对三个叶片的机械耦合力矩')

# ================================================================
#              参数表
# ================================================================
add_section('附录：系统参数表')

add_param_table(
    ['参数', '符号', '数值', '单位'],
    [
        ['定子电阻', 'R', '0.5', 'Ω'],
        ['电感(d/q)', 'L_s', '5', 'mH'],
        ['永磁磁链', 'ψ_f', '0.175', 'Wb'],
        ['极对数', 'p', '4', '-'],
        ['电机惯量', 'J_m', '0.001', 'kg·m²'],
        ['摩擦系数', 'B', '0.001', 'N·m·s/rad'],
        ['传动比', 'N', '1000', '-'],
        ['传动效率', 'η', '0.95', '-'],
        ['齿轮间隙', 'Δ', '0.1', '°'],
        ['风轮半径', 'R_r', '63', 'm'],
        ['空气密度', 'ρ', '1.225', 'kg/m³'],
        ['额定风速', 'V_rated', '12', 'm/s'],
        ['转矩常数', 'K_t', '1.05', 'N·m/A'],
        ['预测时域', 'N_p', '20', '-'],
        ['控制时域', 'N_c', '5', '-'],
        ['采样周期', 'T_s', '0.1', 's'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('K_t = (3/2) × p × ψ_f = 1.5 × 4 × 0.175 = 1.05 N·m/A')
run.italic = True
run.font.color.rgb = RGBColor(100, 100, 100)

# ========== 保存 ==========
import sys
sys.stdout.reconfigure(encoding='utf-8')
output_path = r'D:\mpc控制变桨系统\simulink_pitch\双驱变桨MPC系统公式汇总.docx'
doc.save(output_path)
print(f'文件已保存: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')

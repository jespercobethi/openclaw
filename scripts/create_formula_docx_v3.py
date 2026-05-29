"""
生成双驱变桨MPC系统公式汇总 docx 文件（LaTeX版 v3）
所有公式改为 LaTeX 格式
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 标题 ==========
title = doc.add_heading('', level=0)
run = title.add_run('双驱变桨MPC控制系统 — 公式汇总（LaTeX版）')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于模型预测控制的大型风电机组变桨系统协同控制研究\nv3 LaTeX版 — 2026-05-29')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ========== 辅助函数 ==========
def add_section(title_text, level=1):
    h = doc.add_heading(title_text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_subsection(title_text):
    p = doc.add_paragraph()
    run = p.add_run(f'▶ {title_text}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 80, 130)
    return p

def add_formula(label, *args, note=None):
    """添加公式块：标签 + LaTeX公式（可多行） + 可选说明"""
    # 分离公式和说明
    formulas = []
    for a in args:
        formulas.append(a)
    
    p = doc.add_paragraph()
    run = p.add_run(f'{label}：')
    run.bold = True
    run.font.size = Pt(11)
    
    for formula in formulas:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        run2 = p2.add_run(formula)
        run2.font.name = 'Consolas'
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0, 0, 139)
    
    if note:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Cm(1)
        run3 = p3.add_run(f'📌 {note}')
        run3.font.size = Pt(10)
        run3.font.color.rgb = RGBColor(100, 100, 100)
        run3.italic = True

def add_text(text, indent=0, bold=False, size=11):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    return p

def add_param_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

# ================================================================
#   一、Clark 变换
# ================================================================
add_section('一、Clark 变换（ABC → αβ）')

add_text('将三相静止坐标系（A、B、C）变换为两相静止坐标系（α、β），消除三相间的冗余约束（$i_A+i_B+i_C=0$），降低系统维度。', 0.5)

add_formula('Clark变换矩阵（等幅值）',
    r'\begin{bmatrix} i_\alpha \\ i_\beta \end{bmatrix}'
    r' = \frac{2}{3}'
    r'\begin{bmatrix} 1 & -\frac{1}{2} & -\frac{1}{2} \\ 0 & \frac{\sqrt{3}}{2} & -\frac{\sqrt{3}}{2} \end{bmatrix}'
    r'\begin{bmatrix} i_A \\ i_B \\ i_C \end{bmatrix}',
    '取匝比 N₂/N₃ = 2/3，保证变换前后幅值不变。'
    '\n    α轴与A轴重合，β轴超前α轴90°。'
    '\n    适用条件：三相对称系统（$i_A + i_B + i_C = 0$）')

add_formula('Clark反变换',
    r'\begin{bmatrix} i_A \\ i_B \\ i_C \end{bmatrix}'
    r' = \begin{bmatrix} 1 & 0 \\ -\frac{1}{2} & \frac{\sqrt{3}}{2} \\ -\frac{1}{2} & -\frac{\sqrt{3}}{2} \end{bmatrix}'
    r'\begin{bmatrix} i_\alpha \\ i_\beta \end{bmatrix}',
    '由 $i_A + i_B + i_C = 0$ 约束推导，逆矩阵直接写出。'
    '\n    电压和磁链的Clark变换形式与电流完全相同。')

# ================================================================
#   二、Park 变换
# ================================================================
add_section('二、Park 变换（αβ → dq）')

add_text('将两相静止坐标系变换为随转子旋转的dq坐标系，核心目的：将交流量变为直流量，使PI控制器可以无静差跟踪。', 0.5)

add_formula('Park变换矩阵',
    r'\begin{bmatrix} i_d \\ i_q \end{bmatrix}'
    r' = \begin{bmatrix} \cos\theta_e & \sin\theta_e \\ -\sin\theta_e & \cos\theta_e \end{bmatrix}'
    r'\begin{bmatrix} i_\alpha \\ i_\beta \end{bmatrix}',
    '$\\theta_e$ 为转子电角度（d轴与α轴的夹角）。'
    '\n    d轴（direct axis）：与永磁体磁链方向一致。'
    '\n    q轴（quadrature axis）：超前d轴90°电角度。'
    '\n    $\\omega_e = d\\theta_e/dt = p \\cdot \\omega_m$（电角速度 = 极对数 × 机械角速度）')

add_formula('Park反变换',
    r'\begin{bmatrix} i_\alpha \\ i_\beta \end{bmatrix}'
    r' = \begin{bmatrix} \cos\theta_e & -\sin\theta_e \\ \sin\theta_e & \cos\theta_e \end{bmatrix}'
    r'\begin{bmatrix} i_d \\ i_q \end{bmatrix}',
    'Park矩阵为正交矩阵，其逆等于转置。')

add_formula('ABC → dq 完整变换',
    r'\begin{bmatrix} i_d \\ i_q \end{bmatrix}'
    r' = \frac{2}{3}'
    r'\begin{bmatrix} \cos\theta_e & \cos(\theta_e - \frac{2\pi}{3}) & \cos(\theta_e + \frac{2\pi}{3}) \\ -\sin\theta_e & -\sin(\theta_e - \frac{2\pi}{3}) & -\sin(\theta_e + \frac{2\pi}{3}) \end{bmatrix}'
    r'\begin{bmatrix} i_A \\ i_B \\ i_C \end{bmatrix}',
    'Clark + Park 两步合并，一步从三相abc到旋转dq。')

add_text('物理意义：稳态时 $i_d$ 为励磁分量（无功），$i_q$ 为转矩分量（有功）。控制 $i_q$ 即可线性控制电磁转矩 $T_e$。', 0.5)

# ================================================================
#   三、PMSM 数学模型
# ================================================================
add_section('三、PMSM dq 坐标系数学模型')

add_text('永磁同步电机（PMSM）在dq旋转坐标系下的完整数学模型，是后续状态空间建模和MPC设计的基础。', 0.5)

add_subsection('3.1 电压方程')

add_formula('d轴电压方程',
    r'u_d = R \cdot i_d + L_d \frac{di_d}{dt} - \omega_e L_q i_q',
    '物理含义：电阻压降 + 电感压降 + 交叉耦合项（$-\\omega_e L_q i_q$）。'
    '\n    对于表贴式PMSM（$L_d = L_q = L_s$），简化为 $u_d = R i_d + L_s \\frac{di_d}{dt} - \\omega_e L_s i_q$。'
    '\n    交叉耦合项的存在使得d轴和q轴不独立，需要解耦控制。')

add_formula('q轴电压方程',
    r'u_q = R \cdot i_q + L_q \frac{di_q}{dt} + \omega_e L_d i_d + \omega_e \psi_f',
    '物理含义：电阻压降 + 电感压降 + 交叉耦合项（$+\\omega_e L_d i_d$） + 反电动势（$\\omega_e \\psi_f$）。'
    '\n    反电动势 $\\omega_e \\psi_f$ 是电机发电效应的体现，转速越高反电动势越大。'
    '\n    表贴式PMSM简化为 $u_q = R i_q + L_s \\frac{di_q}{dt} + \\omega_e L_s i_d + \\omega_e \\psi_f$。')

add_formula('永磁磁链方程',
    r'\psi_d = L_d \cdot i_d + \psi_f, \quad \psi_q = L_q \cdot i_q',
    '$\\psi_f$ 为永磁体产生的磁链（常数），$\\psi_d$ 和 $\\psi_q$ 为dq轴总磁链。'
    '\n    对于表贴式PMSM，$L_d = L_q = L_s$。')

add_subsection('3.2 电磁转矩')

add_formula('电磁转矩（通用形式）',
    r'T_e = \frac{3}{2} p (\psi_d i_q - \psi_q i_d)'
    r' = \frac{3}{2} p [\psi_f i_q + (L_d - L_q) i_d i_q]',
    '第一项 $\\psi_f i_q$ 为永磁转矩（主分量），第二项为磁阻转矩。'
    '\n    对于表贴式PMSM（$L_d = L_q$），磁阻转矩为零。'
    '\n    id=0 控制下：$T_e = \\frac{3}{2} p \\psi_f i_q = K_t i_q$，线性关系。')

add_formula('转矩常数',
    r'K_t = \frac{3}{2} p \psi_f',
    '$K_t$ 将 q 轴电流 $i_q$ 线性映射为电磁转矩，是矢量控制的核心参数。'
    '\n    典型值：$K_t = 1.5 \\times 4 \\times 0.175 = 1.05$ N·m/A')

add_subsection('3.3 运动方程')

add_formula('旋转运动方程',
    r'J \frac{d\omega_m}{dt} = T_e - T_L - B \omega_m',
    '$J$: 转动惯量（电机转子 + 折算到电机轴的负载惯量）[kg·m²]。'
    '\n    $T_L$: 负载转矩（气动载荷折算到电机轴）[N·m]。'
    '\n    $B$: 粘性摩擦系数 [N·m·s/rad]，通常很小可忽略。'
    '\n    此方程是MPC中转速预测的物理基础。')

add_formula('角速度关系',
    r'\omega_e = p \cdot \omega_m, \quad \omega_m = \frac{d\theta_m}{dt}, \quad \omega_r = \frac{\omega_m}{N}',
    '$\\omega_e$: 电角速度, $\\omega_m$: 电机机械角速度, $\\omega_r$: 风轮角速度, $N$: 齿轮传动比。')

# ================================================================
#   四、状态空间模型
# ================================================================
add_section('四、状态空间模型（MPC建模基础）')

add_text('状态空间模型是MPC控制器的数学基础。需要将连续PMSM模型离散化，才能在数字控制器中实现滚动优化。', 0.5)

add_subsection('4.1 连续状态空间模型')

add_text('状态变量：$\\mathbf{x} = [i_d, i_q, \\omega_m]^T$    输入：$\\mathbf{u} = [u_d, u_q]^T$    扰动：$d = T_L$', 0.5, bold=True)

add_formula('连续状态方程',
    r'\dot{\mathbf{x}} = \mathbf{A}_c \mathbf{x} + \mathbf{B}_c \mathbf{u} + \mathbf{E}_c d',
    '在稳态工作点 $(i_{d0}, i_{q0}, \\omega_{m0})$ 处线性化。')

add_formula('系统矩阵 A_c',
    r'\mathbf{A}_c = \begin{bmatrix} -\frac{R}{L_s} & p\omega_{m0} & 0 \\ -p\omega_{m0} & -\frac{R}{L_s} & 0 \\ 0 & \frac{K_t}{J} & -\frac{B}{J} \end{bmatrix}',
    '$\\mathbf{A}_c$ 中含 $\\omega_m$ 项，原系统是非线性的，线性化后才能用线性MPC。')

add_formula('输入矩阵 B_c',
    r'\mathbf{B}_c = \begin{bmatrix} \frac{1}{L_s} & 0 \\ 0 & \frac{1}{L_s} \\ 0 & 0 \end{bmatrix}',
    '电压输入只影响电流状态，不直接影响转速。')

add_formula('扰动矩阵 E_c',
    r'\mathbf{E}_c = \begin{bmatrix} 0 \\ -\frac{\psi_f p \omega_{m0}}{L_s} \\ -\frac{1}{J} \end{bmatrix}',
    '扰动通过反电动势和负载转矩影响系统。'
    '\n    第二行包含反电动势扰动项 $-\\psi_f p \\omega_{m0}/L_s$。')

add_subsection('4.2 离散化（前向欧拉法）')

add_formula('离散状态方程',
    r'\mathbf{x}(k+1) = \mathbf{A}_d \mathbf{x}(k) + \mathbf{B}_d \mathbf{u}(k) + \mathbf{E}_d d(k)',
    '前向欧拉近似：$\\dot{\\mathbf{x}} \\approx [\\mathbf{x}(k+1) - \\mathbf{x}(k)] / T_s$。'
    '\n    $T_s$ 为采样周期（典型值 0.1s），需满足香农采样定理。'
    '\n    也可以用后向欧拉或双线性变换（Tustin），精度更高但形式更复杂。')

add_formula('离散矩阵',
    r'\mathbf{A}_d = \mathbf{I} + \mathbf{A}_c T_s, \quad'
    r'\mathbf{B}_d = \mathbf{B}_c T_s, \quad'
    r'\mathbf{E}_d = \mathbf{E}_c T_s',
    '前向欧拉法的离散化公式。')

add_formula('离散 A_d 矩阵展开',
    r'\mathbf{A}_d = \begin{bmatrix} 1 - \frac{R T_s}{L_s} & p\omega_{m0} T_s & 0 \\ -p\omega_{m0} T_s & 1 - \frac{R T_s}{L_s} & 0 \\ 0 & \frac{K_t T_s}{J} & 1 - \frac{B T_s}{J} \end{bmatrix}',
    '对角线元素 $1-xxx \\cdot T_s$ 需要满足稳定性条件（所有特征值在单位圆内）。')

add_formula('离散输出方程',
    r'\mathbf{y}(k) = \mathbf{C} \mathbf{x}(k), \quad \mathbf{C} = \begin{bmatrix} 0 & 0 & 1 \\ 0 & 1 & 0 \end{bmatrix}',
    '输出矩阵 $\\mathbf{C}$ 选择需要反馈的状态量：转速 $\\omega_m$ 和电流 $i_q$。')

add_subsection('4.3 MPC 预测方程')

add_formula('N_p 步预测展开',
    r'\mathbf{x}(k+1|k) = \mathbf{A}_d \mathbf{x}(k) + \mathbf{B}_d \mathbf{u}(k) + \mathbf{E}_d d(k)',
    r'\mathbf{x}(k+2|k) = \mathbf{A}_d \mathbf{x}(k+1|k) + \mathbf{B}_d \mathbf{u}(k+1) + \mathbf{E}_d d(k+1)',
    r'\vdots',
    r'\mathbf{x}(k+N_p|k) = \mathbf{A}_d \mathbf{x}(k+N_p-1|k) + \mathbf{B}_d \mathbf{u}(k+N_p-1) + \mathbf{E}_d d(k+N_p-1)',
    '$\\mathbf{x}(k+i|k)$ 表示在时刻 $k$ 对未来 $k+i$ 时刻状态的预测。'
    '\n    预测依赖：当前状态 $\\mathbf{x}(k)$、未来控制序列 $\\mathbf{U}$、未来扰动估计 $\\mathbf{D}$。'
    '\n    风速预测前馈的作用：提供更准确的 $d(k+i)$ 估计值。')

add_formula('紧凑预测矩阵形式',
    r'\mathbf{X} = \boldsymbol{\Psi} \mathbf{x}(k) + \boldsymbol{\Theta} \mathbf{U} + \boldsymbol{\Gamma} \mathbf{D}',
    r'\mathbf{X} = [\mathbf{x}(k+1), \mathbf{x}(k+2), \ldots, \mathbf{x}(k+N_p)]^T',
    r'\mathbf{U} = [\mathbf{u}(k), \mathbf{u}(k+1), \ldots, \mathbf{u}(k+N_c-1)]^T',
    r'\mathbf{D} = [d(k), d(k+1), \ldots, d(k+N_p-1)]^T',
    '$\\boldsymbol{\\Psi}$: 自由响应矩阵（由 $\\mathbf{A}_d$ 的幂次组成）。'
    '\n    $\\boldsymbol{\\Theta}$: 强制响应矩阵（由 $\\mathbf{A}_d$ 和 $\\mathbf{B}_d$ 组成）。'
    '\n    $\\boldsymbol{\\Gamma}$: 扰动响应矩阵（由 $\\mathbf{A}_d$ 和 $\\mathbf{E}_d$ 组成）。')

# ================================================================
#   五、MPC 目标函数与约束
# ================================================================
add_section('五、MPC 目标函数与约束处理')

add_text('MPC的核心思想：在每个采样时刻，求解一个有限时域优化问题，只执行最优控制的第一步，然后重新测量、重新优化（滚动优化）。', 0.5)

add_subsection('5.1 目标函数')

add_formula('MPC 目标函数（含同步误差）',
    r'J = \sum_{i=1}^{N_p} Q \|\theta_{\text{pitch}}(k+i) - \theta_{\text{ref}}(k+i)\|^2'
    r' + \sum_{j=0}^{N_c-1} R \|\Delta \mathbf{u}(k+j)\|^2'
    r' + \sum_{i=1}^{N_p} S \|\theta_1(k+i) - \theta_2(k+i)\|^2',
    '第1项（跟踪）：桨距角跟踪参考值，$Q$ 越大跟踪越快但可能振荡。'
    '\n    第2项（平滑）：抑制控制量剧烈变化，保护执行机构，$R$ 越大越平滑。'
    '\n    第3项（同步）：两台电机桨距角之差，$S$ 越大同步性越好（核心创新点）。'
    '\n    三者的权重平衡是MPC调参的关键。')

add_formula('控制增量定义',
    r'\Delta \mathbf{u}(k) = \mathbf{u}(k) - \mathbf{u}(k-1)',
    '控制增量用于惩罚控制量的剧烈变化，实现平滑控制。')

add_formula('权重参数选择指南',
    r'Q \text{（跟踪权重）：起步取 } 1 \sim 10 \text{，逐步增大直到跟踪误差满足要求。}',
    r'R \text{（平滑权重）：起步取 } 0.1 \sim 1 \text{，逐步增大直到控制量变化率在执行机构允许范围内。}',
    r'S \text{（同步权重）：起步取 } 5 \sim 20 \text{，确保同步误差小于允许值（如 } 0.1°\text{）。}',
    '调参顺序：先调 $Q$ 和 $R$（单电机跟踪），再加 $S$（双电机同步）。'
    '\n    $Q/R$ 比值决定响应速度 vs 平滑性的权衡。'
    '\n    $S/Q$ 比值决定同步精度 vs 跟踪精度的权衡。')

add_subsection('5.2 约束处理')

add_formula('控制量约束（硬约束）',
    r'u_{d,\min} \leq u_d \leq u_{d,\max}, \quad u_{q,\min} \leq u_q \leq u_{q,\max}',
    '由逆变器直流母线电压决定：$u_{d,\max} \\approx u_{q,\max} \\approx V_{dc}/\\sqrt{3}$。'
    '\n    超出约束范围时，MPC会自动寻找约束边界内的最优解。')

add_formula('状态量约束',
    r'0° \leq \beta \leq 90° \text{（桨距角范围）}',
    r'-10°/\text{s} \leq \frac{d\beta}{dt} \leq 10°/\text{s} \text{（桨距角变化率限制）}',
    r'0 \leq \omega_m \leq \omega_{m,\max} \text{（电机转速上限）}',
    '桨距角变化率限制保护变桨轴承和齿轮箱。'
    '\n    $0°$ 对应满功率发电，$90°$ 对应紧急顺桨（停机保护）。')

add_formula('增量约束',
    r'\Delta \mathbf{u}_{\min} \leq \Delta \mathbf{u}(k) \leq \Delta \mathbf{u}_{\max}',
    '限制每步控制量的最大变化幅度，等效于对控制量变化率的约束。'
    '\n    $\\Delta u_{\\max}$ 由逆变器开关频率和电机电流变化率决定。')

add_formula('QP 标准形式',
    r'\min_{\mathbf{U}} \quad \frac{1}{2} \mathbf{U}^T \mathbf{H} \mathbf{U} + \mathbf{f}^T \mathbf{U}',
    r'\text{s.t.} \quad \mathbf{A}_{\text{ineq}} \mathbf{U} \leq \mathbf{b}_{\text{ineq}}, \quad \mathbf{U}_{\min} \leq \mathbf{U} \leq \mathbf{U}_{\max}',
    '$\\mathbf{H} = \\boldsymbol{\\Theta}^T \\tilde{\\mathbf{Q}} \\boldsymbol{\\Theta} + \\tilde{\\mathbf{R}}$ （海森矩阵，正定保证凸优化）。'
    '\n    $\\mathbf{f} = \\boldsymbol{\\Theta}^T \\tilde{\\mathbf{Q}} (\\boldsymbol{\\Psi} \\mathbf{x}(k) - \\mathbf{X}_{\\text{ref}}) + \\ldots$ （梯度向量）。'
    '\n    每个采样时刻在线求解一次QP，可用 quadprog / MPC Toolbox。')

# ================================================================
#   六、变桨传动模型
# ================================================================
add_section('六、变桨传动模型')

add_text('电机通过减速齿轮箱驱动叶片旋转，传动比通常为1000:1左右。齿轮箱引入了间隙、摩擦等非线性。', 0.5)

add_formula('桨距角与电机角度关系',
    r'\beta = \frac{\theta_m}{N} \ [\text{rad}], \quad \beta = \frac{\theta_m}{N} \times \frac{180}{\pi} \ [°]',
    '$\\theta_m$: 电机转子角度 [rad]，$N$: 传动比（典型值 1000:1）。'
    '\n    大传动比意味着电机转很多圈，叶片才转一点点角度。')

add_formula('转矩折算',
    r"T_L' = \frac{T_{\text{aero}}}{N \cdot \eta}",
    '$T_{\\text{aero}}$: 叶片上的气动转矩 [N·m]，$\\eta$: 传动效率（典型值 0.95）。'
    '\n    大传动比折算后，电机轴上的负载转矩大幅减小（除以N）。')

add_formula('齿轮间隙非线性（迟滞模型）',
    r'\text{if } |\theta_{\text{in}} - \theta_{\text{out}}| < \Delta: \quad \frac{d\theta_{\text{out}}}{dt} = 0 \quad \text{（死区）}',
    r'\text{else: } \quad \theta_{\text{out}} = \theta_{\text{in}} - \Delta \cdot \text{sign}\left(\frac{d\theta_{\text{in}}}{dt}\right) \quad \text{（正常传动）}',
    '$\\Delta$: 齿轮间隙（典型值 0.1°）。'
    '\n    间隙会导致跟踪误差和极限环振荡，是变桨系统的主要非线性源之一。'
    '\n    MPC可以通过约束来部分补偿间隙的影响。')

add_formula('等效惯量（折算到电机轴）',
    r'J_{\text{eq}} = J_m + \frac{J_L}{N^2}',
    '$J_m$: 电机转子惯量，$J_L$: 负载（叶片）惯量。'
    '\n    大传动比使得负载惯量折算到电机轴后大幅缩小（除以$N^2$）。')

# ================================================================
#   七、气动载荷模型
# ================================================================
add_section('七、气动载荷模型')

add_text('风轮从风中捕获的气动能量是变桨系统的负载来源，也是MPC需要处理的主要扰动。', 0.5)

add_formula('风轮气动转矩',
    r'T_{\text{aero}} = \frac{1}{2} \rho \pi R^3 \frac{C_p(\beta, \lambda)}{\lambda} V^2',
    '$\\rho = 1.225$ kg/m³（标准大气密度）。'
    '\n    $R$: 风轮半径 [m]（5MW风机典型值 63m）。'
    '\n    $V$: 来流风速 [m/s]。'
    '\n    $C_p(\\beta, \\lambda)$: 风能利用系数（无量纲，理论最大值 $16/27 \\approx 0.593$，Betz极限）。')

add_formula('叶尖速比',
    r'\lambda = \frac{\omega_r R}{V}',
    '$\\omega_r$: 风轮角速度 [rad/s]。'
    '\n    $\\lambda$ 反映叶片尖端线速度与风速之比，典型最优值 6~8。'
    '\n    $\\lambda$ 过小：叶片失速，效率低；$\\lambda$ 过大：阻力增大，效率也低。')

add_formula('Cp 经验公式',
    r'C_p(\beta, \lambda) = 0.5176 \left( \frac{116}{\lambda_i} - 0.4\beta - 5 \right) e^{-21/\lambda_i} + 0.0068\lambda',
    r'\frac{1}{\lambda_i} = \frac{1}{\lambda + 0.08\beta} - \frac{0.035}{\beta^3 + 1}',
    '$\\lambda_i$ 为等效叶尖速比，修正了桨距角的影响。'
    '\n    $\\beta = 0°$ 时 $C_p$ 最大（约 0.48），$\\beta$ 增大 $C_p$ 急剧下降。'
    '\n    变桨控制的本质：通过调节 $\\beta$ 来调节 $C_p$，从而控制气动功率。')

add_formula('气动功率',
    r'P_{\text{aero}} = \frac{1}{2} \rho \pi R^2 C_p(\beta, \lambda) V^3',
    '功率与风速的三次方成正比，风速小幅变化会引起功率大幅波动。'
    '\n    额定风速以上通过增大 $\\beta$ 来降低 $C_p$，维持额定功率。')

add_formula('气动推力（轴向力）',
    r'F_{\text{thrust}} = \frac{1}{2} \rho \pi R^2 C_t(\beta, \lambda) V^2',
    '$C_t$ 为推力系数，$C_t = 4a(1-a)$，$a$ 为轴向诱导因子。'
    '\n    推力影响塔架载荷，是结构设计的重要输入。')

add_formula('气动转矩折算到电机轴',
    r'T_L = \frac{T_{\text{aero}}}{N \cdot \eta} = \frac{1}{2} \rho \pi R^3 \frac{C_p}{\lambda N \eta} V^2',
    '最终作为状态空间模型中的扰动项 $d(k)$ 输入MPC。')

# ================================================================
#   八、矢量控制与解耦
# ================================================================
add_section('八、矢量控制与解耦（id = 0）')

add_text('矢量控制（FOC）是PMSM的标准控制策略，通过将三相电流分解为d轴和q轴分量，实现转矩的线性控制。id=0控制是最常用的策略。', 0.5)

add_subsection('8.1 前馈解耦')

add_formula('解耦电压方程',
    r"u_d = u_d' - \omega_e L_s i_q, \quad u_q = u_q' + \omega_e L_s i_d + \omega_e \psi_f",
    "$u_d', u_q'$ 为PI控制器的输出（不含耦合项）。"
    '\n    加入前馈补偿后，d轴和q轴完全解耦，可独立设计PI参数。'
    '\n    不解耦时：d轴和q轴电流会互相干扰，导致控制性能下降。')

add_subsection('8.2 电流环 PI 设计')

add_formula('电流环 PI 控制器',
    r"u_d' = K_{p,i} (i_{d,\text{ref}} - i_d) + K_{i,i} \int (i_{d,\text{ref}} - i_d) dt",
    r"u_q' = K_{p,i} (i_{q,\text{ref}} - i_q) + K_{i,i} \int (i_{q,\text{ref}} - i_q) dt",
    '$K_{p,i}$: 电流环比例增益，$K_{i,i}$: 电流环积分增益。'
    '\n    典型设计：将电流环整定为一阶系统，带宽 $f_{\\text{bw},i} \\approx 1$ kHz。'
    '\n    $K_{p,i} = L_s \\omega_c$，$K_{i,i} = R \\omega_c$，其中 $\\omega_c = 2\\pi f_{\\text{bw},i}$。')

add_subsection('8.3 速度环 PI 设计')

add_formula('速度环 PI 控制器',
    r'i_{q,\text{ref}} = K_{p,\omega} (\omega_{\text{ref}} - \omega_m) + K_{i,\omega} \int (\omega_{\text{ref}} - \omega_m) dt',
    '速度环输出作为q轴电流参考值。'
    '\n    速度环带宽应远小于电流环带宽（典型 1/10），避免相互干扰。'
    '\n    $K_{p,\\omega} = J \\omega_{c,\\omega} / K_t$，$K_{i,\\omega} = J \\omega_{c,\\omega}^2 / (K_t N_{\\text{speed}})$。')

add_formula('id=0 控制下的转矩关系',
    r'T_e = K_t \cdot i_q \quad \text{（线性关系）}',
    '当 $i_d=0$ 时，所有电流都用于产生转矩，效率最高。'
    '\n    这是矢量控制的理论基础：控制 $i_q$ 等效于控制转矩。')

# ================================================================
#   九、SVPWM 逆变器模型
# ================================================================
add_section('九、SVPWM 逆变器模型')

add_text('空间矢量脉宽调制（SVPWM）是将控制器输出的电压指令转换为逆变器开关信号的方法。在Simulink建模中需要用到。', 0.5)

add_formula('逆变器输出电压范围',
    r'u_{d,\max} = u_{q,\max} = \frac{V_{dc}}{\sqrt{3}}',
    '$V_{dc}$: 直流母线电压（典型值 540V 或 600V）。'
    '\n    SVPWM 的线性调制区最大输出电压为 $V_{dc}/\\sqrt{3}$。'
    '\n    超出此范围进入过调制区，输出电压波形失真。')

add_formula('电压矢量合成',
    r'\mathbf{V}_{\text{ref}} = \frac{2}{3} (V_a + V_b e^{j2\pi/3} + V_c e^{j4\pi/3})',
    'SVPWM 将参考电压矢量 $\\mathbf{V}_{\\text{ref}}$ 分解为相邻两个基本矢量的加权合成。'
    '\n    一个PWM周期内切换3个基本矢量（含零矢量），实现平滑输出。')

add_formula('占空比计算',
    r'T_1 = \frac{\sqrt{3} T_s}{V_{dc}} \left[ V_\alpha \sin\left(\frac{\pi}{3} - \theta\right) - V_\beta \cos\left(\frac{\pi}{3} - \theta\right) \right]',
    r'T_2 = \frac{\sqrt{3} T_s}{V_{dc}} \left[ -V_\alpha \sin\theta + V_\beta \cos\theta \right]',
    r'T_0 = T_s - T_1 - T_2',
    '$T_1, T_2$: 两个有效矢量的作用时间，$T_0$: 零矢量作用时间。'
    '\n    $\\theta$ 为参考电压矢量的角度。')

# ================================================================
#   十、MPC 滚动优化流程
# ================================================================
add_section('十、MPC 滚动优化流程')

add_formula('MPC 算法步骤',
    r'\text{① 测量/估计当前状态 } \mathbf{x}(k)',
    r'\text{② 读取风速预测值 } d(k) \sim d(k+N_p-1) \text{（前馈）}',
    r'\text{③ 求解 QP 优化问题} \rightarrow \text{得到最优控制序列 } \mathbf{U}^*',
    r'\text{④ 只执行第一步 } \mathbf{u}^*(k)',
    r'\text{⑤ } k \leftarrow k+1 \text{，返回步骤①}',
    '$N_p = 20$（预测时域），$N_c = 5$（控制时域）。'
    '\n    $N_p > N_c$ 意味着只优化前 $N_c$ 步控制量，后面保持不变。'
    '\n    每步都重新求解，实现"滚动"优化，具有天然的反馈校正能力。')

# ================================================================
#   十一、双电机同步控制
# ================================================================
add_section('十一、双电机同步控制')

add_text('双驱变桨系统中，每片叶片由2台电机协同驱动。同步控制的目标是使两台电机的桨距角一致，避免力矩分配不均。', 0.5)

add_formula('同步误差定义',
    r'e_{\text{sync}} = \theta_1 - \theta_2',
    '$\\theta_1, \\theta_2$ 为两台电机的桨距角。'
    '\n    $e_{\\text{sync}} = 0$ 表示两台电机完全同步。'
    '\n    MPC目标函数中的 $S \\|e_{\\text{sync}}\\|^2$ 项驱动同步误差趋近于零。')

add_formula('偏差耦合补偿（传统方法）',
    r'u_{1,\text{comp}} = K_c (\theta_1 - \theta_2), \quad u_{2,\text{comp}} = K_c (\theta_2 - \theta_1)',
    '$K_c$: 耦合增益。偏差耦合是传统多电机同步的经典方法。'
    '\n    本项目将同步误差直接嵌入MPC目标函数，无需额外补偿器。')

add_formula('双电机状态向量',
    r'\mathbf{X} = [i_{d1}, i_{q1}, \omega_{m1}, \theta_1, i_{d2}, i_{q2}, \omega_{m2}, \theta_2]^T',
    r'\dot{\mathbf{X}} = \mathbf{A}_{\text{sys}} \mathbf{X} + \mathbf{B}_{\text{sys}} \mathbf{U} + \mathbf{E}_{\text{sys}} \mathbf{D} + \mathbf{C}_{\text{couple}}',
    '$\\mathbf{A}_{\\text{sys}}$ 为 $8 \\times 8$ 系统矩阵（两台电机的状态拼接）。'
    '\n    $\\mathbf{C}_{\\text{couple}}$: 机械耦合力矩（两台电机通过齿轮箱驱动同一叶片）。'
    '\n    MPC可以直接对 8 维状态空间设计控制器，统一处理跟踪和同步。')

add_formula('三叶片完整系统',
    r'\mathbf{X}_{\text{total}} = [\mathbf{X}_{\text{blade1}}, \mathbf{X}_{\text{blade2}}, \mathbf{X}_{\text{blade3}}]^T \quad \text{（24维状态向量）}',
    '每个叶片 2 台电机 $\\times$ 4 个状态 = 8 维，3 个叶片共 24 维。'
    '\n    实际工程中可以简化：3个叶片的控制问题可以解耦为3个独立的双电机问题。'
    '\n    因为3个叶片之间的耦合主要是气动耦合（轮毂），相对较弱。')

# ================================================================
#   十二、LSTM 风速预测
# ================================================================
add_section('十二、LSTM 风速预测（第三阶段）')

add_text('风速是变桨系统的主要扰动源。通过LSTM神经网络预测未来风速，可以作为MPC的前馈信息，提前做出控制响应。', 0.5)

add_formula('气动转矩预测公式',
    r'T_{L,\text{pred}} = \frac{1}{2} \rho A C_p(\theta, \lambda) \frac{v_{\text{pred}}^3}{\omega}',
    '$v_{\\text{pred}}$: LSTM预测的风速。'
    '\n    将预测风速代入气动公式，得到预测气动转矩，作为MPC的已知扰动 $d(k+i)$。')

add_formula('Kaimal 湍流风谱',
    r'S(f) = \frac{4 \sigma^2 L}{\left(1 + \frac{6fL}{V_{\text{hub}}}\right)^{1.5}} \cdot \frac{1}{f}',
    '$\\sigma$: 湍流标准差 [m/s]，与湍流强度 $I_{\\text{turb}}$ 和平均风速有关：$\\sigma = I_{\\text{turb}} \\cdot V_{\\text{hub}}$。'
    '\n    $L$: 湍流长度尺度（轮毂高度处典型值 340m）。'
    '\n    $V_{\\text{hub}}$: 轮毂高度平均风速 [m/s]。'
    '\n    此公式用于生成仿真用的湍流风速时序数据。')

add_text('LSTM 模型结构：2层LSTM(隐藏层64) + 1层全连接', 0.5)
add_text('输入：过去30s风速序列（采样周期0.1s，共300个点）', 0.5)
add_text('输出：未来30s风速预测值', 0.5)

# ================================================================
#   十三、功率与效率
# ================================================================
add_section('十三、功率与效率方程')

add_formula('电磁功率',
    r'P_e = T_e \cdot \omega_m = \frac{3}{2} (u_d i_d + u_q i_q) \quad \text{[忽略电阻损耗]}',
    '电磁功率等于电磁转矩乘以机械角速度。')

add_formula('输入电功率',
    r'P_{\text{in}} = \frac{3}{2} (u_d i_d + u_q i_q)',
    'abc坐标系和dq坐标系下的功率计算等价（功率不变变换）。')

add_formula('机械功率与损耗',
    r'P_{\text{mech}} = P_e - P_{\text{loss}}, \quad P_{\text{loss}} = \frac{3}{2} R (i_d^2 + i_q^2) + B \omega_m^2',
    '$P_{\\text{loss}}$ 包括铜耗（电阻损耗）和机械摩擦损耗。')

add_formula('风轮捕获功率',
    r'P_{\text{wind}} = \frac{1}{2} \rho \pi R^2 V^3, \quad P_{\text{aero}} = C_p \cdot P_{\text{wind}}',
    '$P_{\\text{wind}}$ 为风的动能功率，$P_{\\text{aero}}$ 为风轮实际捕获功率。'
    '\n    $C_p$ 最大约 0.48（远低于Betz极限 0.593），实际运行点通常在 0.3~0.45。')

# ================================================================
#   十四、风速分布
# ================================================================
add_section('十四、风速统计分布')

add_formula('Weibull 分布（风速概率密度）',
    r'f(V) = \frac{k}{c} \left(\frac{V}{c}\right)^{k-1} \exp\left[-\left(\frac{V}{c}\right)^k\right]',
    '$k$: 形状参数（典型值 2，此时退化为Rayleigh分布）。'
    '\n    $c$: 尺度参数，与平均风速有关：$c = V_{\\text{mean}} / \\Gamma(1+1/k)$。'
    '\n    用于评估风电场年发电量和变桨系统工作时间分布。')

add_formula('Rayleigh 分布（k=2 特例）',
    r'f(V) = \frac{2V}{c^2} \exp\left[-\left(\frac{V}{c}\right)^2\right], \quad c = \frac{2 V_{\text{mean}}}{\sqrt{\pi}}',
    '最常用的风速分布模型，适用于大多数风电场。')

# ================================================================
#   参数表
# ================================================================
add_section('附录A：系统参数表')

add_param_table(
    ['参数', '符号', '数值', '单位', '备注'],
    [
        ['定子电阻', '$R$', '0.5', 'Ω', 'PMSM参数'],
        ['d轴电感', '$L_d$', '5', 'mH', '表贴式 $L_d=L_q$'],
        ['q轴电感', '$L_q$', '5', 'mH', '表贴式 $L_d=L_q$'],
        ['永磁磁链', '$\\psi_f$', '0.175', 'Wb', '永磁体产生'],
        ['极对数', '$p$', '4', '-', '电机极对数'],
        ['电机惯量', '$J_m$', '0.001', 'kg·m²', '转子转动惯量'],
        ['摩擦系数', '$B$', '0.001', 'N·m·s/rad', '粘性摩擦'],
        ['转矩常数', '$K_t$', '1.05', 'N·m/A', '$K_t=\\frac{3}{2}p\\psi_f$'],
        ['传动比', '$N$', '1000', '-', '齿轮箱减速比'],
        ['传动效率', '$\\eta$', '0.95', '-', '齿轮箱效率'],
        ['齿轮间隙', '$\\Delta$', '0.1', '°', '间隙非线性'],
        ['风轮半径', '$R_r$', '63', 'm', '5MW风机'],
        ['空气密度', '$\\rho$', '1.225', 'kg/m³', '标准大气'],
        ['额定风速', '$V_{\\text{rated}}$', '12', 'm/s', '切入额定风速'],
        ['预测时域', '$N_p$', '20', '-', 'MPC预测步数'],
        ['控制时域', '$N_c$', '5', '-', 'MPC控制优化步数'],
        ['采样周期', '$T_s$', '0.1', 's', '控制周期100ms'],
        ['直流母线电压', '$V_{dc}$', '540', 'V', '逆变器输入'],
        ['电流环带宽', '$f_{\\text{bw},i}$', '1000', 'Hz', '电流环设计参数'],
        ['速度环带宽', '$f_{\\text{bw},\\omega}$', '100', 'Hz', '速度环设计参数'],
    ]
)

doc.add_paragraph()
add_text('转矩常数计算：$K_t = \\frac{3}{2} \\times p \\times \\psi_f = 1.5 \\times 4 \\times 0.175 = 1.05$ N·m/A', 0, size=10)

# ================================================================
#   附录B：公式速查表
# ================================================================
add_section('附录B：公式速查表')

add_param_table(
    ['类别', 'LaTeX 公式', '用途'],
    [
        ['Clark变换', '$i_\\alpha, i_\\beta = f(i_A, i_B, i_C)$', '三相→两相静止'],
        ['Park变换', '$i_d, i_q = f(i_\\alpha, i_\\beta, \\theta_e)$', '两相静止→旋转'],
        ['电压方程', '$u_{d/q} = Ri + L\\frac{di}{dt} + \\text{耦合项}$', 'PMSM电气模型'],
        ['电磁转矩', '$T_e = K_t \\cdot i_q$ (id=0)', '转矩-电流关系'],
        ['运动方程', '$J\\frac{d\\omega}{dt} = T_e - T_L - B\\omega$', '转速动态'],
        ['状态空间', '$\\mathbf{x}(k+1) = \\mathbf{A}_d\\mathbf{x}(k) + \\mathbf{B}_d\\mathbf{u}(k) + \\mathbf{E}_d d(k)$', '离散MPC模型'],
        ['MPC目标', '$J = Q\\cdot\\text{跟踪}^2 + R\\cdot\\text{平滑}^2 + S\\cdot\\text{同步}^2$', '核心优化目标'],
        ['气动转矩', '$T_{\\text{aero}} = \\frac{1}{2}\\rho\\pi R^3 \\frac{C_p}{\\lambda} V^2$', '风轮载荷'],
        ['Cp公式', '$C_p = f(\\beta, \\lambda)$ 经验公式', '风能利用系数'],
        ['叶尖速比', '$\\lambda = \\omega_r R / V$', '无量纲转速'],
        ['前馈解耦', '$u_d = u_d\' - \\omega_e L i_q$', 'd/q轴解耦'],
        ['同步误差', '$e_{\\text{sync}} = \\theta_1 - \\theta_2$', '双电机同步'],
        ['SVPWM', '$u_{\\max} = V_{dc}/\\sqrt{3}$', '逆变器电压极限'],
        ['Weibull', '$f(V) = \\frac{k}{c}(\\frac{V}{c})^{k-1} e^{-(V/c)^k}$', '风速分布'],
    ]
)

# ========== 保存 ==========
output_path = r'D:\mpc控制变桨系统\simulink_pitch\双驱变桨MPC系统公式汇总.docx'
doc.save(output_path)
print(f'文件已保存: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading_cn(doc, text, level=1, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if level == 0:
        run.font.size = Pt(22)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(13)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text, bold=False, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_ref(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(2)
    return p

# ==================== 封面 ====================
doc.add_paragraph()
doc.add_paragraph()
add_heading_cn(doc, '附件6', level=2, bold=False)
doc.add_paragraph()
add_heading_cn(doc, '2026年湖南工程学院研究生科研创新项目', level=0)
add_heading_cn(doc, '申请书', level=0)
doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    '项目名称：基于多电机驱动的大型风电变桨伺服系统研究',
    '主 持 人：庄云龙',
    '所在学院：研究生院',
    '联合企业：（若不是联合申报项目则不填）',
    '所属学科/专业类别：能源动力',
    '联系电话：17364504623',
    '电子信箱：2279409386@qq.com',
    '申请日期：',
]
for line in info_lines:
    add_body(doc, line, indent=False)

doc.add_paragraph()
doc.add_paragraph()
add_body(doc, '湖南工程学院研究生院（研究生工作部）', indent=False)
add_body(doc, '2026年', indent=False)

doc.add_page_break()

# ==================== 二、立项依据 ====================
add_heading_cn(doc, '二、立项依据', level=1)

add_heading_cn(doc, '1. 研究目的', level=2)

add_body(doc, '风能作为应用最广泛和发展最快的新能源发电技术，在全球能源转型中占据重要战略地位。根据全球风能理事会（GWEC）发布的《全球风能报告2025》，截至2024年底，全球风电累计装机容量已突破1,100GW，其中我国风电累计装机容量超过470GW，连续多年位居世界第一。在"碳达峰、碳中和"目标的驱动下，我国风电产业正经历从规模化扩张向高质量发展的深刻转变，风电机组单机容量持续增大，陆上机组已普遍达到6-8MW级，海上机组更是向15MW级以上迈进。')

add_body(doc, '变桨伺服系统是大型风力发电机组的核心子系统之一，其主要功能是通过调节叶片桨距角来控制风轮吸收的气动功率，实现机组在额定风速以上的恒功率运行，同时在紧急情况下将叶片顺桨至安全位置以保护机组安全。变桨系统的性能直接关系到风电机组的发电效率、载荷水平和运行安全。随着风电机组单机容量向10MW乃至15MW级迈进，叶片长度超过100米，转动惯量和所需变桨力矩急剧增大，传统由单台电机驱动单个叶片的变桨方案面临严峻挑战：一方面，单台电机所需输出力矩急剧升高，导致电机体积和重量大幅增加，难以适配轮毂内部狭小的安装空间；另一方面，单电机方案缺乏冗余能力，一旦电机或驱动器故障，将无法完成顺桨动作，严重威胁机组安全。为解决上述问题，国外西门子歌美飒、维斯塔斯等主流风机厂商已率先采用双驱电动变桨技术，即每片叶片由2台电机协同驱动同一套变桨机构，实现力矩均分、减小单机容量、提高系统冗余度与运行可靠性[1]。在此架构下，3片叶片共需6台变桨电机，每片叶片的2台电机需高精度同步控制以确保力矩均衡，同时3片叶片之间也需协调一致以保证风轮气动载荷对称。')

add_body(doc, '然而，在实际运行中，双驱变桨系统面临多层级同步控制的挑战：在叶片层级，同一叶片的2台电机由于电气参数差异、减速齿轮箱传动间隙不同等因素，输出力矩难以完全均衡，导致叶片内部产生附加应力；在风轮层级，3片叶片承受的气动载荷不均衡，桨距角难以完全同步。桨距角不同步将导致风轮气动载荷的不对称分布，增加轮毂、主轴和塔架的疲劳载荷，严重时甚至可能引发叶片根部裂纹、主轴承过早失效等结构性损伤。据统计，变桨系统故障是导致风电机组停机的主要原因之一，约占总故障停机时间的20%以上。因此，研究面向双驱（多电机驱动）变桨系统的高性能协同控制技术，既要实现同一叶片内多台电机的力矩均衡同步，又要保证多叶片之间的桨距角协调一致，对于提高大型风电机组运行可靠性、降低运维成本、延长使用寿命具有重要意义。')

add_body(doc, '此外，风能具有较强的随机性和间歇性，风速和风向在短时间内可能发生显著变化。传统变桨控制策略基于当前测量的风速和功率偏差进行被动调节，存在不可避免的响应滞后问题。若能利用先进的预测算法实现风速的短期预测（10-60秒），并将预测结果以前馈方式融入变桨控制回路，将有效提升变桨系统的前瞻性和响应速度，实现从"被动响应"到"主动预判"的转变。')

add_body(doc, '基于上述分析，本项目提出"基于多电机驱动的大型风电变桨伺服系统研究"，以双驱变桨架构为基础（每片叶片由2台电机协同驱动），将模型预测控制（MPC）、多电机协同控制和风速预测前馈三者有机结合，旨在解决叶片内多电机力矩同步控制问题、叶片间桨距角协调控制问题和变桨响应滞后问题，为大型风电机组的安全高效运行提供技术支撑。')

# 2. 国内外研究现状
add_heading_cn(doc, '2. 国内外研究现状', level=2)

add_heading_cn(doc, '2.1 变桨控制技术研究现状', level=3)

add_body(doc, '变桨控制技术经历了从定桨距失速控制到变桨距主动控制的发展过程。在控制算法方面，传统PID控制因结构简单、参数调节方便而被广泛采用，其控制原理清晰，工程实现门槛较低，在早期风电机组控制中发挥了重要作用。然而，面对风电机组强非线性、时变参数以及复杂来流风况等特性，PID控制因参数固定难以自适应调整，其适应性和鲁棒性存在明显不足。近年来，模型预测控制（MPC）因其可显式处理系统约束、具有滚动优化和反馈校正机制等优势，在风电变桨领域受到越来越多的关注。MPC通过构建预测模型并在有限时域内滚动求解优化问题，能够有效兼顾控制性能与约束满足，在处理多变量耦合系统时表现出独特的优越性。然而，目前MPC在变桨控制中的应用多集中于单电机场景，即针对单个变桨执行机构进行独立控制，而将其扩展到多电机驱动场景的研究尚不充分，存在较大的探索空间。')

add_heading_cn(doc, '2.2 多电机协同控制研究现状', level=3)

add_body(doc, '目前主流的多电机同步控制策略主要包括主从控制、偏差耦合控制和环形耦合控制三种。主从控制采用串级结构，将主动电机的输出作为从动电机的参考值，控制结构简单易于实现，但同步精度较低，当从动电机受到扰动时响应滞后明显。偏差耦合控制通过计算任意两电机之间的速度偏差并实施补偿，显著提高了同步精度，但计算复杂度随电机数量呈平方关系增长，不利于大规模推广应用。环形耦合控制则在相邻电机之间建立耦合关系，在保持较好同步性能的同时有效降低了计算复杂度，非常适合多电机场景。在风电双驱变桨领域，每片叶片由2台电机协同驱动，需要在叶片层级实现高精度力矩同步，同时在风轮层级实现3片叶片的桨距角协调。李奔[2]针对海上大型风机双驱变桨系统，采用改进偏差耦合结构结合自抗扰控制（ADRC）实现了双电机同步控制，并通过仿真验证了该策略在负载不平衡工况下的有效性，但其控制策略未考虑系统约束优化，也未涉及预测控制方法。将MPC的预测优化能力与多电机同步控制策略进行深度融合，应用于双驱变桨系统的研究尚属空白，具有重要的理论探索价值。')

add_heading_cn(doc, '2.3 风速预测技术研究现状', level=3)

add_body(doc, '对于变桨控制系统而言，最具应用价值的是超短期风速预测，其预测时域通常为10秒至60秒，能够为变桨动作提供前瞻性参考。传统时间序列方法如自回归积分滑动平均模型（ARIMA）计算简单、理论成熟，但建模能力有限，难以捕捉风速序列中的非线性特征和长程依赖关系。近年来，深度学习方法异军突起，其中长短期记忆网络（LSTM）通过引入门控机制解决了传统循环神经网络中的梯度消失问题，在风速预测任务中展现出强大潜力。然而，将风速预测结果以前馈方式引入变桨控制回路的研究尚处于起步阶段，预测误差的传播机理与补偿策略仍需深入研究。')

add_body(doc, '综合来看，目前将MPC、多电机协同控制和风速预测三者融合应用于双驱（多电机驱动）变桨系统的研究尚属空白，本项目具有明确的创新空间和较高的研究价值。')

# 3. 应用前景
add_heading_cn(doc, '3. 项目应用前景与学术价值', level=2)

add_body(doc, '本项目研究成果可直接应用于大型风力发电机组的变桨伺服系统优化。在工程应用层面，双驱变桨多电机协同控制技术可有效降低叶片内多电机力矩同步误差和叶片间桨距角同步误差，减小风轮不对称载荷，延长机组关键部件的使用寿命，为海上大型风电装备的可靠运行提供技术支撑。在学术研究层面，本项目将MPC理论与多电机同步控制理论进行融合创新，构建了"预测-优化-同步-前馈"的完整控制框架，丰富了风电变桨控制的理论体系，为相关领域的后续研究提供了可借鉴的理论基础与方法论参考。')

# 4. 研究基础
add_heading_cn(doc, '4. 现有研究基础与条件', level=2)

add_body(doc, '本课题组长期从事风电装备设计与制造研究，在风电机组控制领域具有扎实的研究基础，积累了丰富的理论与实验经验。课题组李奔（2026）完成了海上大型风机双驱电动变桨系统的研究，采用改进偏差耦合结构结合ADRC实现了双电机同步控制，并通过Ansys Electronics完成了变桨电机的电磁设计与仿真验证，为本项目的多电机驱动变桨系统建模与控制研究提供了直接的理论与实践基础。此外，王耀锋（2022）完成了基于ADRC+偏差耦合的偏航永磁同步电机多电机同步控制研究，刘世博（2021）完成了风电变桨永磁同步电机无位置传感器控制研究，李晓凤（2019）完成了风电机组变桨用IPMSM的MTPA控制研究。上述前期工作为本项目奠定了坚实的理论与技术基础。课题组拥有高性能计算工作站，已安装MATLAB/Simulink R2024a仿真平台，具备进行大规模系统仿真与算法验证的硬件和软件条件，可保证本项目的顺利推进。')

# 5. 参考文献
add_heading_cn(doc, '5. 参考文献', level=2)

refs = [
    '[1] 穆安乐,张惠明,刘小虎,等.大型风电机组变桨距控制技术研究综述[J].电力系统自动化,2018,42(21):25-35.',
    '[2] 李奔.海上大型风机双驱电动变桨系统研究[D].湖南工程学院,2026.',
    '[3] 王明军,李梅,张伟.基于模型预测控制的风电变桨系统研究[J].太阳能学报,2020,41(8):256-263.',
    '[4] 刘军,郭庆安,李强.风电变桨系统PID参数优化策略研究[J].电机与控制学报,2019,23(6):78-85.',
    '[5] 耿涛,赵强,王海峰.海上风电机组变桨控制技术发展现状与展望[J].船舶工程,2020,42(S1):202-209.',
    '[6] 刘向杰,刘晓枫.基于约束优化的风电MPC变桨控制器设计[J].控制理论与应用,2019,36(4):512-520.',
    '[7] 郝鹏,张国梁,李铭.基于MPC的风力发电机组独立变桨控制[J].中国电机工程学报,2021,41(15):5234-5242.',
    '[8] 蒋为龙,刘宏伟.模型预测控制在风电系统中的应用综述[J].电网技术,2017,41(3):896-904.',
    '[9] 彭川,刘永强.多电机同步控制策略研究综述[J].电机与控制学报,2018,22(9):1-10.',
    '[10] 魏海峰.多电机伺服系统偏差耦合协同控制研究[D].哈尔滨工业大学,2017.',
    '[11] 蒋明,吴青华.偏差耦合控制的多电机同步策略在风电变桨系统中的应用[J].电力系统保护与控制,2020,48(8):169-176.',
    '[12] 刘洋,汪晓健.基于滑模控制的多电机变桨同步策略[J].机械工程学报,2021,57(12):234-243.',
    '[13] 董文永,李青,刘健.基于LSTM的风电场超短期风速预测[J].电力系统自动化,2019,43(7):67-74.',
    '[14] 迟渼洪,张晓东,马最.风电变桨控制技术研究进展与展望[J].中国电机工程学报,2022,42(12):4330-4343.',
    '[15] Corradini M L,Ippoliti G,Orlando G.A model predictive approach for a wind turbine robust pitch control[J].Proceedings of the IEEE Conference on Decision and Control,2010:5376-5381.',
    '[16] van Solingen E,van Wingerden J W.Linear model predictive control for wind turbine pitch control during partial load region[J].IFAC Proceedings Volumes,2011,44(1):3716-3721.',
    '[17] J,Jonsson M,Umesh K N.Robust MPC for wind turbine pitch regulation using delta operator parameterization[C].American Control Conference(ACC),2014:3735-3740.',
    '[18] Bossanyi E A.Individual blade pitch control for load reduction[J].Wind Energy,2003,6(2):119-128.',
    '[19] van der Veen G J,van Wingerden J W,Verhaegen M.Global identification of wind turbine rotor impedance using state-space model predictive control[J].Wind Energy,2013,16(7):1003-1021.',
    '[20] Hussain H,Muhtaroglu A,Bader N.Model predictive pitch control of wind turbine using LIDAR-based wind speed measurement[J].Renewable Energy,2019,135:1342-1351.',
    '[21] Du Y,Wu J,Li S.Model predictive control for wind turbine pitch angle based on online optimization[J].Energy,2020,198:117384.',
    '[22] Liu Y,Chen Z,Li C.Disturbance observer-based MPC for wind turbine pitch control with input constraints[J].Renewable Energy,2021,175:731-740.',
    '[23] Shi J,Li Z,Li J.Improved ring coupling control for multi-motor servo systems[J].IEEE Transactions on Industrial Electronics,2019,66(12):9219-9228.',
    '[24] Perez T,Fossen T I.Singular perturbation analysis of a wind turbine multi-motor pitch control system[J].Ocean Engineering,2018,159:432-445.',
    '[25] Wang S,Lv B,Liu Y,et al.Application of dual-drive pitch technology in offshore wind power[J].Ship Engineering,2019,41(S1):288-290+294.',
    '[26] Wang J,Wang Y,Li Y.Ultra-short term wind speed forecasting using LSTM neural network with feature selection[J].Renewable Energy,2020,151:1248-1256.',
    '[27] Liu H,Mi X,Li Y.Wind speed prediction method using deep learning with temporal convolutional network[J].Energy,2021,225:120197.',
    '[28] Chen K,Yu J.Short-term wind speed prediction using EEMD-based LSTM neural network[J].Energy,2019,174:973-987.',
    '[29] Hu J,Zheng W,Wang C.A combined model based on CEEMDAN and LSTM for wind speed prediction[J].Energy Reports,2022,8:10347-10357.',
    '[30] Yang T,Wang Y,Li W.Feedforward compensation control for wind turbine pitch using wind speed prediction[J].Wind Energy,2021,24(8):783-798.',
    '[31] Xu H,Cheng Q,Wang J.A comprehensive review of wind turbine pitch control strategies and coordinated optimization methods[J].Renewable and Sustainable Energy Reviews,2023,173:113082.',
]

for ref in refs:
    add_ref(doc, ref)

doc.add_page_break()

# ==================== 三、研究方案 ====================
add_heading_cn(doc, '三、研究方案', level=1)

add_heading_cn(doc, '3.1 研究目标、研究内容和拟解决的关键问题', level=2)

add_heading_cn(doc, '3.1.1 研究目标', level=3)

add_body(doc, '设计一种基于模型预测控制（MPC）的双驱变桨多电机协同伺服控制系统，以每片叶片2台电机协同驱动的双驱架构为基础，实现叶片内双电机的高精度力矩同步控制以及叶片间桨距角的协调控制，并通过风速预测前馈提升变桨系统的前瞻响应能力。具体目标包括：叶片内双电机力矩同步误差较传统PI控制降低50%以上；叶片间桨距角同步误差较传统PI控制降低50%以上；功率波动标准差较无前馈方案降低30%以上；变桨机构疲劳动作频率降低20%以上。')

add_heading_cn(doc, '3.1.2 研究内容', level=3)

add_body(doc, '（1）多电机驱动变桨伺服系统建模。建立永磁同步电机（PMSM）的数学模型，包括电压方程、磁链方程、电磁转矩方程和运动方程，采用矢量控制（id=0）策略实现电机转矩的线性解耦控制。在此基础上，构建包含减速齿轮箱传动比、传动效率和间隙特性的单电机变桨伺服系统模型。进一步，考虑同一叶片内2台电机通过变桨机构的机械耦合效应，建立叶片内双电机耦合动力学模型。最终，将3片叶片（共6台电机）的模型组合，构建完整的多电机驱动变桨系统动力学模型，参考李奔[2]所采用的22kW变桨PMSM参数及传动机构参数。')

add_body(doc, '（2）MPC变桨控制器设计。以PMSM的状态空间模型为基础，设计有限时域MPC控制器。优化目标函数包含桨距角跟踪误差的加权平方和与控制量增量的加权平方和。约束条件包括桨距角变化速率约束（通常为±10°/s）和电机转矩约束。通过求解二次规划（QP）问题获得最优控制序列。')

add_body(doc, '（3）多电机协同控制策略研究。针对双驱变桨系统的层级结构（叶片内2电机同步+风轮3叶片协调），设计多层次协同控制方案。在叶片层级，基于偏差耦合或环形耦合结构实现同一叶片内2台电机的力矩同步控制；在风轮层级，协调3片叶片的桨距角一致性。重点研究将叶片内同步误差和叶片间协调误差同时纳入MPC优化目标函数的一体化设计方法，实现变桨跟踪、叶片内同步和叶片间协调的统一优化。')

add_body(doc, '（4）风速预测前馈设计。基于LSTM神经网络建立风速超短期预测模型，以历史风速数据为输入，预测未来10~60秒的风速变化趋势。将预测风速转换为预测风轮气动转矩，作为已知扰动输入MPC预测模型。')

add_body(doc, '（5）仿真验证与性能评估。在MATLAB/Simulink环境下搭建完整仿真平台，设计多种仿真工况，将本项目所提融合控制方法与传统PI控制、单一MPC控制进行对比，定量评估叶片内力矩同步误差、叶片间桨距角同步误差、功率波动、变桨动作频率等性能指标。')

add_heading_cn(doc, '3.1.3 拟解决的关键问题', level=3)

add_body(doc, '（1）多电机驱动变桨系统的多层次耦合动力学建模问题（叶片内双电机机械耦合+风轮级气动耦合）；')
add_body(doc, '（2）MPC控制器中预测模型精度与计算实时性的平衡问题；')
add_body(doc, '（3）风速预测误差对前馈控制效果的影响及鲁棒性补偿问题。')

add_heading_cn(doc, '3.2 拟采取的研究方法及可行性分析', level=2)

add_heading_cn(doc, '3.2.1 研究方法', level=3)

add_body(doc, '理论建模法：基于永磁同步电机的Park方程和风能利用系数理论，建立多电机驱动变桨伺服系统的状态空间模型。')
add_body(doc, '仿真分析法：利用MATLAB/Simulink搭建系统仿真模型，通过参数扫描和灵敏度分析研究关键参数对控制性能的影响。')
add_body(doc, '对比实验法：设计统一的仿真测试平台和评价指标体系，将本项目所提方法与传统方法进行严格对比。')
add_body(doc, '文献研究法：系统梳理国内外最新研究成果，明确创新切入点。')

add_heading_cn(doc, '3.2.2 可行性分析', level=3)

add_body(doc, '理论可行性方面，MPC理论经过40余年的发展已相当成熟，多电机同步控制理论在数控机床等领域已有广泛应用，课题组李奔（2026）已在双驱变桨系统同步控制方面取得了初步成果。技术可行性方面，课题组已系统掌握MATLAB/Simulink建模仿真技术，具备Python编程能力。条件可行性方面，课题组拥有高性能计算工作站，已安装MATLAB R2024a，可满足全部仿真计算需求。')

add_heading_cn(doc, '3.3 本项目的创新之处', level=2)

add_body(doc, '（1）MPC与双驱变桨多电机同步控制的一体化融合设计。不同于传统"独立控制器+同步补偿器"的分离式架构，本项目针对双驱变桨系统的层级结构特点，将叶片内双电机力矩同步误差和叶片间桨距角协调误差同时纳入MPC优化目标函数，在统一的优化框架下实现桨距角跟踪、叶片内同步和叶片间协调的多目标优化。')
add_body(doc, '（2）风速预测前馈与MPC的深度融合。在MPC预测模型中引入LSTM预测的未来风速信息作为已知扰动，使控制器具备"提前预判"能力。不同于传统前馈控制的静态补偿，MPC框架下的前馈补偿是基于滚动优化的动态补偿。')
add_body(doc, '（3）面向双驱变桨系统的多层次同步优化策略。针对每片叶片2台电机的对称结构特点，在叶片层级采用改进偏差耦合或环形耦合实现力矩均衡，在风轮层级通过MPC协调3片叶片的桨距角一致性，形成"叶片内同步—叶片间协调"的多层次控制架构。')

add_heading_cn(doc, '3.4 预期研究进展', level=2)

add_body(doc, '第一阶段（2026年6月—2026年8月）：文献调研与理论学习；完成PMSM变桨伺服系统建模；搭建基础Simulink仿真平台。')
add_body(doc, '第二阶段（2026年9月—2026年11月）：建立叶片内双电机耦合动力学模型及风轮级6电机系统模型；设计MPC变桨控制器；完成多层次协同控制方案设计。')
add_body(doc, '第三阶段（2026年12月—2027年2月）：训练LSTM风速预测模型；将风速预测前馈融入MPC框架；完成融合控制算法联调。')
add_body(doc, '第四阶段（2027年3月—2027年5月）：多工况仿真验证与性能对比分析；撰写学术论文并投稿。')
add_body(doc, '第五阶段（2027年6月）：撰写结题报告；整理提交全部成果材料。')

add_heading_cn(doc, '3.5 预期成果', level=2)

add_body(doc, '（1）发表学术论文1~2篇，其中SCI或中文核心期刊论文至少1篇；')
add_body(doc, '（2）申请发明专利1项（双驱变桨多电机协同控制方法相关）；')
add_body(doc, '（3）完成研究报告1份（含完整仿真模型和技术文档）；')
add_body(doc, '（4）建立可复用的双驱变桨伺服系统MATLAB/Simulink仿真模型1套。')

# 保存
out_path = r'C:\Users\xuan1\.openclaw\workspace\main\基于多电机驱动的大型风电变桨伺服系统研究-修改版.docx'
doc.save(out_path)
print(f'Done: {out_path}')

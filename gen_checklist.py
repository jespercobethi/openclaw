from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
title = doc.add_heading('', level=0)
run = title.add_run('基于多电机驱动的大型风电变桨伺服系统研究\n——研究内容执行清单')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# ===== 一、多电机驱动变桨系统动力学建模 =====
h1 = doc.add_heading('一、多电机驱动变桨系统动力学建模', level=1)

doc.add_heading('1.1 需要搞定的事', level=2)

items_1 = [
    '整理PMSM的dq坐标系数学模型（电压方程、转矩方程、机械方程）',
    '确定电机参数（22kW IPMSM，参考李晓凤2019论文：Rs, Ld, Lq, ψf, p, J, B）',
    '建立双电机耦合传动模型（齿轮啮合刚度kg、阻尼cg、传动比）',
    '建立风载荷模型（Kaimal湍流谱 + 阵风模型）',
    '推导气动阻力矩公式（Cp查表或拟合，ρ, A, R等参数）',
    '整理完整状态空间方程（8状态：id1, iq1, ω1, θ1, id2, iq2, ω2, θ2）',
    '在工作点处线性化',
    '离散化（前向欧拉，Ts=0.1s），得到Ad, Bd矩阵',
]
for item in items_1:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.2 需要的资料', level=2)
refs_1 = [
    '李晓凤(2019)论文 — 电机参数',
    '王耀锋(2022)论文 — 双电机耦合结构参考',
    '李奔(2026)论文 — 风机变桨机构参数',
]
for r in refs_1:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('1.3 需要的工具', level=2)
tools_1 = [
    'MATLAB（符号计算推导公式用syms）',
    'Simulink（搭开环模型）',
]
for t in tools_1:
    doc.add_paragraph(t, style='List Bullet')

# ===== 二、基于MPC的多电机协同控制器设计 =====
h2 = doc.add_heading('二、基于MPC的多电机协同控制器设计', level=1)

doc.add_heading('2.1 需要搞定的事', level=2)

items_2 = [
    '学透MPC理论（状态空间→预测方程→QP求解流程）',
    '推导预测方程（从Ad, Bd推到H, f矩阵）',
    '设计目标函数权重矩阵Q, R, S的初始值',
    '确定同步误差权重qsync的初始值',
    '编写MATLAB MPC求解代码（quadprog或MPC Toolbox）',
    '实现约束处理（力矩、转速、角速率、桨距角范围）',
    '调通单电机MPC → 再扩展到双电机',
    '整热启动策略（用上一时刻解作为初始猜测）',
    '测试单步求解时间，确保 < 100ms',
]
for item in items_2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.2 需要的资料', level=2)
refs_2 = [
    'Rawlings《Model Predictive Control》教材',
    'B站DR_CAN的MPC教程（入门用）',
    'MATLAB MPC Toolbox文档',
]
for r in refs_2:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('2.3 需要的工具', level=2)
tools_2 = [
    'MATLAB MPC Toolbox（mpc, mpcmove命令）',
    'quadprog函数（二次规划求解器）',
]
for t in tools_2:
    doc.add_paragraph(t, style='List Bullet')

# ===== 三、风速预测与前馈补偿 =====
h3 = doc.add_heading('三、风速预测与前馈补偿', level=1)

doc.add_heading('3.1 需要搞定的事', level=2)

items_3 = [
    '用Kaimal湍流模型生成风速数据集（1000条，不同均值/湍流强度）',
    '搭建LSTM网络（PyTorch：2层LSTM + 1层FC）',
    '准备训练数据（输入过去30s风速，输出未来10s预测）',
    '训练LSTM模型，调参（学习率、隐藏单元数）',
    '验证预测精度（RMSE目标 < 0.5 m/s）',
    '把预测风速转为预测气动转矩TL_hat',
    '修改MPC预测模型，加入扰动输入项Bw·TL_hat',
    '实现误差衰减加权（远端预测权重低）',
    '设计扰动观测器补偿预测误差（可选，后期加）',
]
for item in items_3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 需要的工具', level=2)
tools_3 = [
    'Python + PyTorch（训练LSTM）',
    'MATLAB（生成风速数据、与MPC对接）',
    '训练用GPU（没有的话CPU也能跑，数据量不大）',
]
for t in tools_3:
    doc.add_paragraph(t, style='List Bullet')

# ===== 四、仿真验证与对比分析 =====
h4 = doc.add_heading('四、仿真验证与对比分析', level=1)

doc.add_heading('4.1 需要搞定的事', level=2)

items_4 = [
    '搭建完整Simulink仿真平台（风速模块 + 双电机系统 + 控制器）',
    '实现ADRC+偏差耦合控制器（对标王耀锋方案）',
    '实现PID+主从控制器（传统方案基线）',
    '跑6个工况 × 3种方法 = 18组仿真',
    '记录数据：跟踪误差RMSE、最大同步误差、调节时间、超调量、功率波动标准差',
    '做敏感性分析（Np、qsync、预测误差、Ts各4组 = 16组仿真）',
    '画对比图表（跟踪曲线、同步误差曲线、柱状图对比）',
    '撰写对比分析报告',
]
for item in items_4:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 工况清单', level=2)

# 工况表格
table = doc.add_table(rows=7, cols=3, style='Table Grid')
table.cell(0, 0).text = '工况编号'
table.cell(0, 1).text = '工况名称'
table.cell(0, 2).text = '风速条件'

cases = [
    ('1', '额定风速稳态', '恒定v=12m/s'),
    ('2', '阵风冲击', '12→17→12 m/s阶跃突变'),
    ('3', '湍流连续扰动', '均值12m/s + Kaimal湍流（强度15%）'),
    ('4', '单电机故障降级', 't=5s时电机1力矩降至50%'),
    ('5', '参数不确定性', 'Rs, Ld, Lq偏差±20%'),
    ('6', '风速预测误差', '预测误差±15%'),
]
for i, (num, name, cond) in enumerate(cases):
    table.cell(i+1, 0).text = num
    table.cell(i+1, 1).text = name
    table.cell(i+1, 2).text = cond

doc.add_paragraph('')

doc.add_heading('4.3 需要的工具', level=2)
tools_4 = [
    'Simulink',
    'MATLAB画图（plot, bar, subplot）',
]
for t in tools_4:
    doc.add_paragraph(t, style='List Bullet')

# ===== 五、建议执行顺序 =====
h5 = doc.add_heading('五、建议执行顺序', level=1)

steps = [
    ('第一步', '学MPC理论', '2周'),
    ('第二步', '推导PMSM模型 + 线性化离散化', '2周'),
    ('第三步', '写单电机MPC代码跑通', '2周'),
    ('第四步', '扩展到双电机 + 加同步误差项', '2周'),
    ('第五步', '生成风速数据 + 训练LSTM', '2周（和第四步并行）'),
    ('第六步', '前馈融入MPC + 全系统联调', '2周'),
    ('第七步', '跑对比实验 + 敏感性分析', '3周'),
    ('第八步', '整理数据 + 画图', '1周'),
]

table2 = doc.add_table(rows=9, cols=3, style='Table Grid')
table2.cell(0, 0).text = '阶段'
table2.cell(0, 1).text = '内容'
table2.cell(0, 2).text = '预计时间'
for i, (step, content, time) in enumerate(steps):
    table2.cell(i+1, 0).text = step
    table2.cell(i+1, 1).text = content
    table2.cell(i+1, 2).text = time

doc.add_paragraph('')
p_total = doc.add_paragraph()
run = p_total.add_run('总计约15-16周，前面几个月主要花在学理论和调代码上。')
run.bold = True

# 保存
output_path = r'C:\Users\xuan1\.openclaw\workspace\main\MPC风电变桨研究-执行清单.docx'
doc.save(output_path)
print(f'已保存到: {output_path}')
